from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DATE_LABEL = "15 July 2026"
AUTHOR_LABEL = "LERU Integrity Map project team"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "222222"
MUTED = "5B6570"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F4F6F9"
MID_GRAY = "D8DEE6"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120


CATEGORY_LABELS = {
    "local-output": "Institution-owned local output",
    "national-or-sector-output": "National, regional or sector output",
    "procedure-only": "Procedure-visible, output-light",
    "restricted-or-internal-output": "Restricted or internal output signal",
    "historical-or-case-specific": "Historical or case-specific output",
}


SYSTEM_GROUPS = [
    (
        "Belgium",
        "Federal container with community-level lanes",
        [
            {
                "institution": "KU Leuven",
                "category": "local-output",
                "route": "Commission on Research Integrity, Research Integrity Reporting Desk and Research Integrity Advisers. Informal concerns can go to advisers; formal allegations enter through the reporting and commission route.",
                "output": "Annual reports provide counts, categories and anonymized advice summaries. They are among the clearest institution-owned accountability signals in the LERU sample.",
                "angle": "Use KU Leuven to show how a local reporting corridor can sit inside a wider Flemish and Belgian system. The next evidence task is row-level indexing by year, matter type, outcome and publication exclusion.",
            }
        ],
    ),
    (
        "Denmark",
        "National scientific-dishonesty route plus local questionable-practice handling",
        [
            {
                "institution": "University of Copenhagen",
                "category": "local-output",
                "route": "Practice Committee and Named Persons, alongside the national Danish Board on Research Misconduct for the statutory scientific-dishonesty route.",
                "output": "The 2024 Practice Committee annual report is the strongest local institutional publication example identified in the Danish pass.",
                "angle": "Use Copenhagen to explain a two-level allocation model: national handling for the statutory core and local handling for returned cases and other questionable practices. Index annual reports and minutes by case type and outcome.",
            }
        ],
    ),
    (
        "Finland",
        "National TENK guideline and statement route with institutional first-instance responsibility",
        [
            {
                "institution": "University of Helsinki",
                "category": "national-or-sector-output",
                "route": "Research-integrity adviser and support route, followed by a written notification to the Chancellor; TENK remains the national second-line statement route after local handling.",
                "output": "No Helsinki-owned annual integrity statistics or local case archive was located. Public visibility sits mainly in national TENK statement summaries and annual reporting.",
                "angle": "Use Helsinki to show why national output should not be relabelled as institutional output. Member validation should test whether any non-personal local statistics or Helsinki-specific statement links are public.",
            }
        ],
    ),
    (
        "France",
        "Legally codified, institutionally executed and nationally coordinated RIS network",
        [
            {
                "institution": "Universite Paris-Saclay",
                "category": "procedure-only",
                "route": "POLETHIS and a network of scientific-integrity referents; the CER-PS ethics committee remains a research-ethics boundary rather than a misconduct-output body.",
                "output": "No public institutional misconduct archive, annual RIS activity report, signalement statistics page or anonymized decision channel was identified.",
                "angle": "Use Paris-Saclay as an output-light but route-visible example. Validate whether annual RIS activity or non-personal case-learning material exists beyond the public route pages.",
            },
            {
                "institution": "Sorbonne University",
                "category": "restricted-or-internal-output",
                "route": "Scientific-integrity delegation, RIS/RIO, integrity committee and ambassador network.",
                "output": "A visible annual statement or reporting source exists, but member validation indicates that the studied case-level reporting is internal or restricted-access rather than a public case archive.",
                "angle": "Use Sorbonne to distinguish visible governance reporting from openly reusable case output. Retain the restricted/internal category unless a public case table or case-summary channel is identified.",
            },
            {
                "institution": "University of Strasbourg",
                "category": "national-or-sector-output",
                "route": "Scientific-integrity referent who verifies reports, may open an inquiry, can seek expert advice and submits recommendations to the university president.",
                "output": "No Strasbourg-owned annual RIS report or anonymized case channel was located. The Ofis synthesis is national context, not local Strasbourg output.",
                "angle": "Use Strasbourg to show how a well-described institutional route can remain locally output-light while national coordination is visible. Validate local statistics or case-learning material.",
            },
        ],
    ),
    (
        "Germany",
        "DFG good-research-practice baseline with institution-specific ombuds and commission routes",
        [
            {
                "institution": "University of Freiburg",
                "category": "historical-or-case-specific",
                "route": "Ombudsperson or self-control route, Investigation Commission and a coordination office for honesty in science.",
                "output": "No current standing annual output channel was identified. Freiburg does maintain a substantial historical public corridor around the sports-medicine and doping misconduct complex.",
                "angle": "Use Freiburg to distinguish case-complex transparency from routine accountability. Validate whether current annual ombuds or commission statistics exist.",
            },
            {
                "institution": "Heidelberg University",
                "category": "procedure-only",
                "route": "Ombudsmen for good academic practice and the Commission for the Safeguarding of Good Academic Practice.",
                "output": "The rules require general anonymized reporting to the Rector, but no public annual version, statistics page or decision archive was located.",
                "angle": "Use Heidelberg to examine the public/private boundary around internal anonymized governance reporting. Member validation should ask whether any non-personal annual output is published.",
            },
            {
                "institution": "Ludwig-Maximilians-Universitat Munchen / LMU Munich",
                "category": "procedure-only",
                "route": "Ombudspersons for good scientific practice and an investigation committee under the 2023 LMU good-scientific-practice regulation.",
                "output": "No public case-output channel, annual ombudsperson statistics page or anonymized decision archive was found; public communication appears discretionary after a final decision.",
                "angle": "Use LMU to discuss discretionary publication rules. Locate a current office or committee page and retest rectorate reporting for non-personal summaries.",
            },
        ],
    ),
    (
        "Ireland",
        "NRIF and IUA coordination with institutional policies and national aggregate reporting",
        [
            {
                "institution": "Trinity College Dublin",
                "category": "national-or-sector-output",
                "route": "Senior Dean as College Research Integrity Officer, with allegation management delegated to the Dean of Research and upheld matters routed to the relevant disciplinary process.",
                "output": "No Trinity-owned annual statement, local aggregate statistics or case-learning channel was located. NRIF reports provide national aggregate context without Trinity-specific counts.",
                "angle": "Use Trinity to show the limits of inferring local transparency from a national aggregate series. Validate whether non-personal local reporting exists outside the public Senior Dean and research pages.",
            }
        ],
    ),
    (
        "Italy",
        "Distributed institutional system with ethics and integrity responsibilities that need careful boundary separation",
        [
            {
                "institution": "University of Milan",
                "category": "restricted-or-internal-output",
                "route": "University Ethics Committee supported by the Research Ethics Office, with a mixed project-ethics and non-disciplinary code-opinion role.",
                "output": "No public misconduct archive or aggregate integrity-statistics channel was identified. Committee minutes and opinions appear restricted, while annual reporting to the Academic Senate is not publicly located.",
                "angle": "Use Milan to demonstrate why project ethics, code opinions and misconduct handling cannot be merged. Validate the publication status of the committee president's annual report and any non-personal opinions.",
            }
        ],
    ),
    (
        "Netherlands",
        "Institution-first complaint handling under the national code, with LOWI as a second-line advisory route",
        [
            {
                "institution": "University of Amsterdam",
                "category": "local-output",
                "route": "Academic Integrity Committee (CWI) and the institutional complaints-regulation route.",
                "output": "The UvA CWI annual report 2024 provides advice summaries and yearly complaint and advice counts.",
                "angle": "Use UvA to compare annual-report fields with Leiden and Utrecht. Extract the 2024 fields and test for earlier or later reports.",
            },
            {
                "institution": "Leiden University",
                "category": "local-output",
                "route": "Academic Integrity Committee with separate Leiden University and LUMC chambers.",
                "output": "Leiden publishes multi-year CWI annual reports and an overview of advice and final Executive Board judgments, including anonymized case rows.",
                "angle": "Use Leiden as a high-value example of local annual reporting plus public judgments. Index by chamber, complaint type, LOWI interaction and final outcome.",
            },
            {
                "institution": "Utrecht University",
                "category": "local-output",
                "route": "Scientific Integrity Committee (CWI), followed where relevant by LOWI and a final Executive Board judgment.",
                "output": "Utrecht's 2024 annual report records case activity, while Universities of the Netherlands publishes anonymized Utrecht case PDFs.",
                "angle": "Use Utrecht to connect institutional annual reporting with sector-hosted case publication. Index allegation type, CWI advice, LOWI step and board judgment.",
            },
        ],
    ),
    (
        "Spain",
        "Distributed system with institutional codes and regional integrity routes",
        [
            {
                "institution": "Universitat de Barcelona",
                "category": "national-or-sector-output",
                "route": "UB research-integrity code and integrity portal, with CIR-CAT as the regional integrity route for the institution-linked recommendation; UB ethics committees remain boundary infrastructure.",
                "output": "No UB-owned annual integrity report or general case archive was located. CIR-CAT Recommendation 1/2025 is a public regional recommendation directly linked to UB doctoral-affiliation guidance.",
                "angle": "Use Barcelona to show how institution-linked regional output differs from institution-owned reporting. Index the recommendation and validate the UB/CIR-CAT division.",
            }
        ],
    ),
    (
        "Sweden",
        "Statutory national handling of research misconduct with local handling of other deviations",
        [
            {
                "institution": "Lund University",
                "category": "local-output",
                "route": "Deviations from Good Research Practice Review Board, Vice-Chancellor route and faculty research representatives, alongside the national Npof route.",
                "output": "Lund publishes an Annual Report 2025 for the local review board. This is local aggregate reporting, not a searchable decision archive.",
                "angle": "Use Lund to explain the Npof/local split and how local annual reporting can make that allocation visible. Index received matters, route, outcome and national/local division.",
            }
        ],
    ),
    (
        "Switzerland",
        "Fragmented cantonal, institutional and funder-linked system",
        [
            {
                "institution": "University of Geneva",
                "category": "procedure-only",
                "route": "Faculty integrity delegates, possible ad hoc fact-finding commission and Rectorate decision route, with ethics committees kept separate.",
                "output": "No annual integrity statistics, ombudsperson report or anonymized decision archive was located. Publication authority is discretionary rather than a standing channel.",
                "angle": "Use Geneva to analyze confidentiality and discretionary publication. Validate whether any non-personal Rectorate summaries exist.",
            },
            {
                "institution": "ETH Zurich",
                "category": "local-output",
                "route": "Scientific Integrity Office, Integrity Commission, Good Scientific Practice Commission and departmental research-integrity advisers.",
                "output": "Official ETH endpoints list anonymized investigation reports and procedure statistics.",
                "angle": "Use ETH as the clearest Swiss institution-owned output corridor, while avoiding generalization to other Swiss universities. Index the reports and statistics tables.",
            },
            {
                "institution": "University of Zurich",
                "category": "procedure-only",
                "route": "Three procedural instances: ombudspersons, Research Integrity Delegate and Executive Board; the Coordination Office supports inquiries and investigations but is not a fourth decision body.",
                "output": "Member validation confirms no current public aggregate statistics, annual integrity report, anonymized decision archive or case-reporting channel. Official news and annual reports are contextual only.",
                "angle": "Use UZH as a member-validated negative-output case and as a version-control example. Recheck the substantially revised ordinance, bodies and URLs when the new framework is expected in early 2027.",
            },
        ],
    ),
    (
        "United Kingdom",
        "Employer-led Concordat model with annual statements as a recurring public-reporting corridor",
        [
            {
                "institution": "University of Cambridge",
                "category": "local-output",
                "route": "Research-integrity reports hub and institutional research-misconduct procedure.",
                "output": "The report hub provides anonymized information on how allegations were assessed and investigated.",
                "angle": "Use Cambridge in the cross-UK annual-statement comparison. Index statement year, allegation stage, investigation count and outcome.",
            },
            {
                "institution": "University of Edinburgh",
                "category": "local-output",
                "route": "Vice Principal Research and Innovation route with delegated College named-person routes and a public annual-statement hub.",
                "output": "Annual research-integrity statements are available from 2014-2015 through 2025-2026 and provide high-level formal-investigation information.",
                "angle": "Use Edinburgh as the longest clearly indexed annual-statement series in the LERU sample. Extract counts, outcomes and lessons-learned depth by year.",
            },
            {
                "institution": "Imperial College London",
                "category": "local-output",
                "route": "Research Misconduct Response Group, Registrar and University Secretary, Vice-Provost and Director of Research Integrity Investigations.",
                "output": "A multi-year annual-statement series is visible. The 2024 statement provides narrative case learning, while the 2025 statement is the latest located report.",
                "angle": "Use Imperial to compare screening, full-investigation and narrative-learning fields, and to track the procedure changes introduced from 2026.",
            },
            {
                "institution": "University College London (UCL)",
                "category": "local-output",
                "route": "Compliance and Assurance route, Research Misconduct Committee analysis and Research, Innovation and Global Engagement Committee oversight.",
                "output": "UCL maintains an annual-statement archive through 2024-2025. The 2023-2024 statement contains long-run committee analysis and tabular allegation and investigation fields.",
                "angle": "Use UCL as a field-rich UK annual-statement case. Extract the 2023-2024 and 2024-2025 records and validate the current direct procedure source.",
            },
            {
                "institution": "University of Oxford",
                "category": "local-output",
                "route": "Research Integrity Statement route, with a Registrar and Proctors' Office distinction for staff and student research matters.",
                "output": "The 2024 statement contains anonymized numbered allegations and outcomes, plus separately described student research-work cases.",
                "angle": "Use Oxford to show both detailed annual-statement reporting and the need to separate student research cases from staff misconduct routes. Index by category, outcome and anonymization depth.",
            },
        ],
    ),
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=CELL_TOP_BOTTOM_DXA, start=CELL_START_END_DXA, bottom=CELL_TOP_BOTTOM_DXA, end=CELL_START_END_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{edge}")
        element = tc_mar.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table, color=MID_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA} DXA: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([r_pr, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc, kind, text_indent_dxa, marker_align_dxa, hanging_dxa, after_twips, line_twips):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(text_indent_dxa))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(text_indent_dxa))
    ind.set(qn("w:hanging"), str(hanging_dxa))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), str(after_twips))
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def style_document(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(DARK_BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle.paragraph_format.keep_with_next = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom = {
        "Kicker": (10, BLUE, True, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 14),
        "Cover Meta": (10.5, MUTED, False, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 4),
        "Lead": (12, NAVY, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 10),
        "Profile Meta": (9.5, MUTED, True, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 6),
        "Table Text": (9.2, INK, False, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0),
        "Table Header": (9.2, NAVY, True, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0),
        "Source Note": (9.5, MUTED, False, True, WD_ALIGN_PARAGRAPH.LEFT, 4, 4),
    }
    for name, (size, color, bold, italic, alignment, before, after) in custom.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15 if name != "Lead" else 1.2


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("LERU Integrity Map | Working report outline")
    set_run_font(left, size=8.5, color=MUTED)
    right = paragraph.add_run(f"\t{DATE_LABEL}")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Working review copy | Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(paragraph)

    first_footer = section.first_page_footer
    paragraph = first_footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Review copy | Track changes enabled")
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc, text, level=1, page_break=False):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    run = paragraph.add_run(text)
    return paragraph


def add_body(doc, text, keep_with_next=False):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_labeled(doc, label, text, keep_with_next=False):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.keep_with_next = keep_with_next
    lead = paragraph.add_run(label + " ")
    set_run_font(lead, size=11, color=NAVY, bold=True)
    value = paragraph.add_run(text)
    set_run_font(value, size=11, color=INK)
    return paragraph


def add_bullet(doc, num_id, text):
    paragraph = doc.add_paragraph(style="Normal")
    apply_num(paragraph, num_id)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_numbered(doc, num_id, title, detail):
    paragraph = doc.add_paragraph(style="Normal")
    apply_num(paragraph, num_id)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first = paragraph.add_run(title + ". ")
    set_run_font(first, size=11, color=NAVY, bold=True)
    second = paragraph.add_run(detail)
    set_run_font(second, size=11, color=INK)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), PALE_BLUE)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    lead = paragraph.add_run(label + " ")
    set_run_font(lead, size=10.5, color=NAVY, bold=True)
    value = paragraph.add_run(text)
    set_run_font(value, size=10.5, color=INK)
    return paragraph


def fill_table_cell(cell, text, header=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Header" if header else "Table Text"
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=9.2, color=NAVY if header else INK, bold=header)


def add_corpus_table(doc):
    rows = [
        ("Country layer", "49 country records", "Comparative European system map"),
        ("Dossier depth", "40 deep dossiers; 9 expanded overview dossiers", "Separates mature dossiers from first expanded profiles"),
        ("Local overview corpus", "40 readable country overview DOCX files", "Country synthesis and extraction baseline"),
        ("Transparency layer", "35 scored jurisdiction entries", "Compares public case and accountability visibility"),
        ("Source registry", "1,566 entries", "Structured audit trail for official web and document sources"),
        ("LERU institution layer", "24 official member profiles", "Institution-level route and public-output comparison"),
        ("LERU strict output flag", "14 profiles", "Verified public-output evidence under the current binary evidence field"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Corpus component", "Current snapshot", "Use in the report"]
    for idx, header in enumerate(headers):
        fill_table_cell(table.rows[0].cells[idx], header, header=True)
        shade_cell(table.rows[0].cells[idx], PALE_BLUE)
    repeat_table_header(table.rows[0])
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            fill_table_cell(cells[idx], value)
    set_fixed_table_geometry(table, [2340, 2520, 4500])
    set_table_borders(table)
    source = doc.add_paragraph(style="Source Note")
    run = source.add_run("Source: current project datasets and logs at the report cut-off; counts should be frozen again before publication.")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    return table


def add_typology_table(doc):
    rows = [
        ("Institution-owned local output", "12", "Annual statements, aggregate tables, anonymized summaries, reports or local decision corridors"),
        ("National, regional or sector output", "4", "Public visibility sits mainly outside the institution's own reporting channel"),
        ("Procedure-visible, output-light", "5", "Local route is public, but no standing local public-output channel was located"),
        ("Restricted or internal output", "2", "Reporting signals exist, but case-level or committee material is not openly accessible"),
        ("Historical or case-specific", "1", "Public material concerns a historical case complex rather than routine current reporting"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Dominant visibility model", "Profiles", "Interpretation"]
    for idx, header in enumerate(headers):
        fill_table_cell(table.rows[0].cells[idx], header, header=True)
        shade_cell(table.rows[0].cells[idx], PALE_BLUE)
    repeat_table_header(table.rows[0])
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            fill_table_cell(cells[idx], value, alignment=WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_fixed_table_geometry(table, [2880, 900, 5580])
    set_table_borders(table)
    source = doc.add_paragraph(style="Source Note")
    run = source.add_run("Interpretation note: these five categories classify dominant public-source visibility across all 24 profiles. They are not a league table and do not measure case incidence or institutional quality.")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    return table


def add_member_matrix(doc):
    records = []
    for country, _, members in SYSTEM_GROUPS:
        for member in members:
            records.append((member["institution"], country, CATEGORY_LABELS[member["category"]]))
    if len(records) != 24:
        raise ValueError(f"Expected 24 LERU members, found {len(records)}")
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(("Institution", "Country", "Dominant public-output category")):
        fill_table_cell(table.rows[0].cells[idx], header, header=True)
        shade_cell(table.rows[0].cells[idx], PALE_BLUE)
    repeat_table_header(table.rows[0])
    for institution, country, category in records:
        cells = table.add_row().cells
        fill_table_cell(cells[0], institution)
        fill_table_cell(cells[1], country)
        fill_table_cell(cells[2], category)
    set_fixed_table_geometry(table, [4320, 1800, 3240])
    set_table_borders(table)
    return table


def add_editorial_table(doc):
    rows = [
        ("1. Freeze scope", "Confirm report versus journal-article format, authorship, audience and data cut-off", "Editorial decision"),
        ("2. Consolidate evidence", "Complete high-yield row indexing and recheck time-sensitive routes", "Evidence refresh"),
        ("3. Member validation", "Invite structured correction of route, office, reporting and boundary descriptions", "Validated institution layer"),
        ("4. Draft synthesis", "Turn section prompts and profile capsules into continuous argument and comparative figures", "Full manuscript draft"),
        ("5. Review and redline", "Use this review copy for tracked revisions, comments and responsibility assignment", "Review-ready draft"),
        ("6. Publish", "Accept revisions, freeze tables, add formal references and prepare accessible PDF/Word outputs", "Final report package"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(("Phase", "Main action", "Output")):
        fill_table_cell(table.rows[0].cells[idx], header, header=True)
        shade_cell(table.rows[0].cells[idx], PALE_BLUE)
    repeat_table_header(table.rows[0])
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            fill_table_cell(cells[idx], value)
    set_fixed_table_geometry(table, [1800, 5400, 2160])
    set_table_borders(table)
    return table


def build_document(output_path: Path):
    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0])
    bullet_num_id = create_numbering(doc, "bullet", 540, 260, 280, 80, 290)
    decimal_num_id = create_numbering(doc, "decimal", 540, 260, 280, 80, 290)

    # Cover: editorial_cover pattern, resolved for a long-form research report.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    kicker = doc.add_paragraph(style="Kicker")
    run = kicker.add_run("LERU INTEGRITY MAP | REPORT DEVELOPMENT")
    set_run_font(run, size=10, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    run = title.add_run("Research Integrity Systems in Europe: Public Visibility, Institutional Routes and the LERU Member Landscape")
    set_run_font(run, size=30, color=NAVY, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    run = subtitle.add_run("A substantive article and report outline based on the full European mapping project")
    set_run_font(run, size=15, color=DARK_BLUE)
    meta = doc.add_paragraph(style="Cover Meta")
    run = meta.add_run("Working report/article outline")
    set_run_font(run, size=10.5, color=MUTED, bold=True)
    meta = doc.add_paragraph(style="Cover Meta")
    run = meta.add_run(f"Prepared by {AUTHOR_LABEL} | {DATE_LABEL}")
    set_run_font(run, size=10.5, color=MUTED)
    meta = doc.add_paragraph(style="Cover Meta")
    run = meta.add_run("Public-source review copy | Not an official LERU statement | Not an institutional audit")
    set_run_font(run, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()

    add_heading(doc, "Executive summary", level=1)
    lead = doc.add_paragraph(style="Lead")
    run = lead.add_run(
        "Europe does not have one research-integrity system. It has a layered field in which law, national bodies, funders, academies, universities and specialist regulators allocate responsibilities differently - and disclose their work with very different levels of public visibility."
    )
    set_run_font(run, size=12, color=NAVY)
    add_body(
        doc,
        "The project covers 49 countries and 40 mature country dossiers. It combines a source-linked static web app with country overviews, committee and case-file directories, transparency coding, source notes and a structured source registry. The report should present this not as a ranking exercise but as a system-mapping study: the central analytical question is where a concern starts, through which route it travels, where a decision ends, and what becomes publicly knowable afterwards.",
    )
    add_body(
        doc,
        "The LERU layer covers all 24 official member universities. Every profile contains identifiable institution-level procedure or route evidence. Under the current strict evidence flag, 14 profiles have verified public-output evidence. A broader typology classifies 12 profiles as institution-owned local output, 4 as national, regional or sector output, 5 as procedure-visible but output-light, 2 as restricted or internal output, and 1 as historical or case-specific. These measures answer different questions and must remain separate in the final text.",
    )
    add_body(
        doc,
        "The report's main claim should be that procedural visibility is becoming a common baseline, while public accountability remains uneven and structurally dependent on national allocation rules, institutional reporting customs and confidentiality choices. Absence of public output is not evidence of absence of cases. Conversely, a visible ethics, clinical-trials, animal-research, data-protection, quality-assurance or IP register should not be counted as general research-misconduct transparency unless the source explicitly connects it to that function.",
    )
    add_body(
        doc,
        "For LERU and its research-integrity network, the project offers a practical route map and a validation instrument. A mature final report could support a shared vocabulary, a minimal public-reporting dataset, recurring route verification and comparative learning without requiring institutions to publish personal, confidential or operationally sensitive case information.",
    )
    add_callout(
        doc,
        "Drafting status.",
        "This is a substantive scaffold rather than a polished final manuscript. Narrative transitions, a formal literature review, figure production, member validation and final source-current checks remain explicit editorial tasks. All figures and counts use the project snapshot dated 15 July 2026.",
    )

    add_heading(doc, "Editorial brief", level=1)
    add_labeled(doc, "Proposed primary audience.", "LERU and INTE members, research-integrity officers, university leaders, funders, national bodies and researchers studying governance and public accountability.")
    add_labeled(doc, "Proposed form.", "A 35-50 page LERU report or white paper, with a shorter journal-article derivative after the report's evidence and terminology have been validated.")
    add_labeled(doc, "Working thesis.", "European research-integrity systems are best compared as layered routes rather than as single national bodies; procedure visibility is widespread, but the public record of handling and outcomes remains heterogeneous.")
    add_labeled(doc, "Tone.", "Analytical, source-linked and non-ranking. The report should be candid about missing evidence and avoid interpreting output visibility as institutional performance or case incidence.")
    add_labeled(doc, "Unit of observation.", "Country system, institutional route, public-output corridor and boundary regime. Each unit should be labelled consistently so adjacent governance does not migrate into the misconduct dataset.")

    add_heading(doc, "Proposed report architecture", level=1)
    architecture = [
        ("Executive summary", "State the central finding, corpus scale, LERU-specific result and practical recommendations in two pages."),
        ("Why map research-integrity systems", "Introduce the fragmentation problem, the public-accountability question and the value of route-based comparison."),
        ("Scope, concepts and methods", "Define research integrity, misconduct, questionable research practices, transparency and boundary regimes; explain source selection and extraction."),
        ("The European system landscape", "Compare national boards, legally anchored networks, code-led systems, academy or funder routes and institution-led models without forcing them into one hierarchy."),
        ("Public accountability and case visibility", "Present a continuum from no national publication through procedure-only visibility, annual summaries and structured archives."),
        ("The LERU member landscape", "Explain the 24-profile methodology, comparative findings and the distinction between strict evidence and the five-category visibility typology."),
        ("Institutional profile capsules", "Cover every LERU member with a concise route, output signal, country context and validation task."),
        ("Cross-case lessons", "Analyze route design, second-line review, publication ownership, annual reporting, confidentiality and boundary discipline."),
        ("Implications and recommendations", "Propose a shared vocabulary, minimum public dataset, route register and member-validation cycle."),
        ("Discussion, limitations and conclusion", "State what the map can and cannot show, then return to the value of system-level comparison."),
    ]
    for title_text, detail in architecture:
        add_numbered(doc, decimal_num_id, title_text, detail)

    add_heading(doc, "Corpus snapshot", level=1)
    add_corpus_table(doc)

    add_heading(doc, "Why map research-integrity systems", level=1, page_break=True)
    add_heading(doc, "Problem statement", level=2)
    add_body(
        doc,
        "European research-integrity arrangements are frequently described through a single prominent actor: a national board, an ombudsperson, a code, a funder rule or a university committee. That shorthand is often misleading. In practice, concerns may start inside an institution, move to a national advisory or adjudicative route, trigger funder consequences, and intersect with employment, disciplinary, data-protection or research-ethics processes. A comparison that records only the most visible body misses the allocation of authority and the path of accountability.",
    )
    add_heading(doc, "Core research questions", level=2)
    questions = [
        "How are responsibility and decision-making distributed across national, funder, academy and institutional actors?",
        "Where does a research-integrity concern enter the system, through which bodies can it travel, and where does the formal decision end?",
        "Which procedures, annual reports, statistics, decisions, opinions or case summaries are publicly visible?",
        "How do LERU member universities translate their country systems into local routes, committees, offices and reporting practices?",
        "Which adjacent regimes must be kept analytically separate to avoid overstating the reach or transparency of general misconduct systems?",
    ]
    for item in questions:
        add_bullet(doc, bullet_num_id, item)
    add_heading(doc, "Contribution", level=2)
    add_body(
        doc,
        "The contribution is both empirical and methodological. Empirically, the project consolidates dispersed official sources into a European route map and a 24-institution LERU layer. Methodologically, it proposes a reproducible distinction between system design, procedure visibility, public output and adjacent governance. The final article should make clear that these are separate variables rather than stages of a single maturity score.",
    )

    add_heading(doc, "Scope, concepts and methods", level=1)
    add_heading(doc, "Scope and evidence strategy", level=2)
    add_body(
        doc,
        "The project uses an official-source-first strategy. National legislation portals, ministries, research-integrity bodies, academies, funders, universities, public research organisations and specialist authorities form the preferred source backbone. Non-official sources are excluded by default or explicitly caveated. Local country folders, overview documents and source notes are checked before new online collection, preserving a visible audit trail.",
    )
    add_body(
        doc,
        "Collection proceeds in stages: initial source discovery, country overview drafting, structured extraction to the app, transparency coding, deep-dossier development, committee and case-file discovery, and quality consolidation. For institutions, the same logic is repeated at a smaller scale: identify the route, procedure, office or committee; test for public reporting; record boundaries and caveats; and formulate member-validation questions.",
    )
    add_heading(doc, "Analytical distinctions", level=2)
    distinctions = [
        "Research misconduct handling versus prospective research-ethics approval.",
        "Scientific-integrity adjudication versus higher-education quality assurance.",
        "Funder monitoring or grant consequences versus a misconduct finding.",
        "Research integrity in research versus academic integrity in teaching and assessment.",
        "Data-protection, IP, whistleblowing and employment routes versus research-integrity procedures.",
        "Public route visibility versus public case or outcome visibility.",
    ]
    for item in distinctions:
        add_bullet(doc, bullet_num_id, item)
    add_heading(doc, "Quality and validation", level=2)
    add_body(
        doc,
        "Quality consolidation should be presented as part of the method rather than as housekeeping. It includes current-version checks, ENRIO cross-checks where available, main-code identification, normalization of committee directories, indexing of the strongest public repositories and explicit treatment of inaccessible or missing sources. In the LERU layer, member validation is a methodological stage: members can correct office names, routes, publication practices and boundary interpretations without requiring disclosure of private case material.",
    )

    add_heading(doc, "The European system landscape", level=1, page_break=True)
    add_heading(doc, "A layered field, not a single model", level=2)
    add_body(
        doc,
        "The European comparison should open with a strong negative proposition: the absence of one national misconduct board does not mean the absence of a system. Some countries rely on statutory national bodies; others combine codes, funders, academies and institutional investigations; federal or devolved systems distribute authority across regions or communities; and several systems place the decisive first instance with employers while using national bodies for advice, appeal or coordination.",
    )
    add_heading(doc, "Proposed comparative typology", level=2)
    typology = [
        "Statutory national adjudication with institution-level intake or complementary handling.",
        "National advisory, appeal or statement body following institutional first-instance investigation.",
        "Legally anchored and nationally coordinated network executed by institutions.",
        "Code-led or self-regulatory system with funder, academy and institutional implementation.",
        "Federal, devolved or community-based system with multiple territorial lanes.",
        "Predominantly institution-led or developing system without a stable national general-misconduct route.",
    ]
    for item in typology:
        add_bullet(doc, bullet_num_id, item)
    add_heading(doc, "System-map narrative", level=2)
    add_body(
        doc,
        "A recurring figure should show four questions: where the network starts, through which bodies a concern can travel, who can make or review a finding, and where public information appears. This route-based figure is more informative than a map that merely places one marker on each country. It also creates a common language for comparing national bodies with LERU institutional routes.",
    )
    add_callout(
        doc,
        "Proposed Figure 1.",
        "A layered route diagram connecting concern intake, institutional assessment, national or sector review, funder consequences and public-output channels. Boundary regimes should appear as adjacent lanes rather than as steps in the general misconduct route.",
    )

    add_heading(doc, "Public accountability and case visibility", level=1)
    add_heading(doc, "The transparency problem", level=2)
    add_body(
        doc,
        "A public procedure answers whether a route exists and how a concern should be handled. It does not necessarily show whether the route is used, what kinds of matters arise, how long they take, what findings are reached or whether institutions learn from them. The transparency layer therefore isolates a narrower question: what public evidence exists about handling and outcomes, and in what format?",
    )
    add_heading(doc, "Proposed publication continuum", level=2)
    continuum = [
        ("No national publication route located", "Only general codes, procedures or unrelated boundary outputs are visible."),
        ("Body or route visible, no standing archive", "The responsible body is identifiable, but routine public output is absent or unclear."),
        ("Summary-based accountability", "Annual reports, aggregate statistics or anonymized narrative summaries provide limited recurring visibility."),
        ("Fragmented or selective publication", "Opinions, statements, press releases or institutional files exist but are dispersed across sites and years."),
        ("Structured archive", "A stable decisions, opinions or statement repository supports repeatable case-level or row-level analysis."),
    ]
    for title_text, detail in continuum:
        add_numbered(doc, decimal_num_id, title_text, detail)
    add_body(
        doc,
        "The project overview currently describes 35 scored jurisdiction entries, including structured archives, fragmented or mixed systems, summary-based or institutional-statement systems, low-publication or procedure-only systems, and one country where no national publication route was identified. Before publication, these totals should be regenerated from the final transparency dataset and accompanied by a clear coding appendix.",
    )
    add_heading(doc, "Interpretive safeguards", level=2)
    safeguards = [
        "Do not infer low case incidence from low public visibility.",
        "Do not interpret a high-output archive as proof of higher misconduct incidence or weaker integrity culture.",
        "Do not combine national, sector and institution-owned output without identifying publication ownership.",
        "Do not count prospective approvals or regulatory registers as general misconduct decisions.",
        "Record anonymization, retention and appeal context before comparing case-level files.",
    ]
    for item in safeguards:
        add_bullet(doc, bullet_num_id, item)

    add_heading(doc, "The LERU member landscape", level=1, page_break=True)
    add_heading(doc, "Coverage and central finding", level=2)
    add_body(
        doc,
        "The institutional layer represents all 24 LERU member universities in the current membership snapshot. Each profile identifies a public procedure, route, office, committee, ombudsperson, adviser network or comparable institutional entry point. The central comparative finding mirrors the country layer: route visibility is more consistent than public-output visibility.",
    )
    add_body(
        doc,
        "Under the current strict evidence field, 14 profiles have verified public-output evidence. The five-category typology below classifies the dominant visibility model for every profile. It is intentionally broader than the binary evidence flag: a profile can have meaningful national, restricted or historical public context without operating an institution-owned standing output channel.",
    )
    add_typology_table(doc)
    add_heading(doc, "Member overview", level=2)
    add_member_matrix(doc)
    add_heading(doc, "Comparative interpretation", level=2)
    insights = [
        "The UK and several Dutch profiles provide the deepest recurring institution-owned reporting corridors in the current sample.",
        "National or sector-level outputs can be highly informative while remaining unsuitable as substitutes for local institutional reporting.",
        "Procedure-only profiles often have clear confidential routes; their output-light classification should not be treated as a performance judgment.",
        "Restricted/internal and historical/case-specific material require separate labels because they provide visibility of a different kind.",
        "Even within the same country, institutions may differ substantially in publication ownership, detail, retention and route presentation.",
    ]
    for item in insights:
        add_bullet(doc, bullet_num_id, item)

    add_heading(doc, "Institutional profile capsules", level=1, page_break=True)
    add_body(
        doc,
        "The capsules below are designed as article-ready building blocks, not final institutional verdicts. Each one identifies the public route, the dominant public-output signal and the analytical role the institution can play in the full report. They should be converted into footnoted prose after member validation and a final source-current check.",
    )
    member_count = 0
    for country, system_context, members in SYSTEM_GROUPS:
        add_heading(doc, f"{country}: {system_context}", level=2)
        for member in members:
            member_count += 1
            add_heading(doc, member["institution"], level=3)
            meta = doc.add_paragraph(style="Profile Meta")
            run = meta.add_run(f"{country} | {CATEGORY_LABELS[member['category']]} | Profile {member_count} of 24")
            set_run_font(run, size=9.5, color=MUTED, bold=True)
            add_labeled(doc, "Visible route.", member["route"], keep_with_next=True)
            add_labeled(doc, "Public-output signal.", member["output"], keep_with_next=True)
            add_labeled(doc, "Role in the report and validation task.", member["angle"])
    if member_count != 24:
        raise ValueError(f"Expected 24 profile capsules, generated {member_count}")

    add_heading(doc, "Cross-case lessons", level=1, page_break=True)
    add_heading(doc, "Route design", level=2)
    add_body(
        doc,
        "The LERU profiles demonstrate that a local route can take several forms: a standing commission, an ombudsperson sequence, a scientific-integrity referent, a reporting desk, an adviser network, an executive decision route or a hybrid. The report should compare the functional sequence - intake, assessment, investigation, advice, decision and review - rather than assuming that identical labels imply identical powers.",
    )
    add_heading(doc, "Publication ownership", level=2)
    add_body(
        doc,
        "A core analytical variable is who publishes. Institution-owned annual statements provide a different accountability signal from national statement summaries, sector-hosted judgments, regional recommendations or funder notices. The final report should label ownership and scope for every output example before comparing content depth.",
    )
    add_heading(doc, "Confidentiality and learning", level=2)
    add_body(
        doc,
        "Confidential handling and public learning are not necessarily opposites. Aggregate counts, anonymized case themes, procedural timelines and lessons-learned notes can provide public accountability without disclosing identifiable case details. The report should examine which institutions already use these middle forms and which safeguards make them credible.",
    )
    add_heading(doc, "Boundary discipline", level=2)
    add_body(
        doc,
        "Institutional sites often make ethics committees, clinical governance, animal research, data protection, open science, IP and quality assurance more visible than misconduct handling. These routes are important for the wider integrity environment, but counting them as misconduct transparency would distort comparison. Boundary discipline is therefore a substantive finding, not merely a coding precaution.",
    )

    add_heading(doc, "Implications and recommendations for LERU and INTE", level=1)
    recommendations = [
        ("Adopt a shared route vocabulary", "Describe intake, informal advice, preliminary assessment, formal investigation, decision, review and publication as distinct functions, even when institutions use different local titles."),
        ("Define a minimal public accountability dataset", "Consider annual publication of non-personal counts, handling stages, broad matter categories, outcome groupings, median or banded duration, review status and lessons learned."),
        ("Maintain a verified route register", "Ask each member to confirm the current public entry page, office or committee name, governing procedure, report location and next planned revision date."),
        ("Separate national and institutional outputs", "Record when a national, regional or sector body supplies the public record and avoid presenting it as institution-owned reporting."),
        ("Keep boundary regimes explicit", "Show ethics, clinical, animal, data, IP, security and quality-assurance routes as adjacent governance unless a source directly integrates them into misconduct handling."),
        ("Use member validation as version control", "Create a light recurring review cycle for corrections, link changes, newly public reports and major procedural reforms."),
    ]
    for title_text, detail in recommendations:
        add_numbered(doc, decimal_num_id, title_text, detail)
    add_callout(
        doc,
        "Proposed Figure 2.",
        "A LERU public-output map that shows the five visibility models without ranking institutions. Each profile should retain its country-system context and publication owner.",
    )

    add_heading(doc, "Discussion and limitations", level=1)
    add_heading(doc, "What the project can show", level=2)
    add_body(
        doc,
        "The map can show publicly documented allocation of responsibility, route design, code and procedure baselines, public-output formats and evidence gaps. It can support structured comparison and targeted validation. It can also reveal where apparently similar systems differ in decision authority or publication ownership.",
    )
    add_heading(doc, "What the project cannot show", level=2)
    limits = [
        "The prevalence of misconduct, because public output and case incidence are not equivalent.",
        "The effectiveness or fairness of a procedure without case-level process and outcome evidence.",
        "The full internal route where institutions appropriately keep operational details non-public.",
        "A definitive current picture without a declared source cut-off and recurring link/version checks.",
        "A single European ranking that remains valid across different legal, institutional and publication models.",
    ]
    for item in limits:
        add_bullet(doc, bullet_num_id, item)
    add_heading(doc, "Sources of uncertainty", level=2)
    add_body(
        doc,
        "Language, search indexing, inaccessible pages, changing URLs and uneven annual-report practices create ascertainment bias. A source not located in one pass may still exist. The final report should therefore distinguish 'none exists' from the more defensible 'no public source was identified in this review' and should date every negative-output claim.",
    )

    add_heading(doc, "Conclusion", level=1)
    add_body(
        doc,
        "A credible European research-integrity comparison must map systems as routes. It should show how responsibilities move between institutions, national bodies, funders and specialist regimes, and it should distinguish a public procedure from a public record of handling and outcomes. The current project provides the evidence architecture for that comparison across 49 country records and a detailed institutional lens across all 24 LERU members.",
    )
    add_body(
        doc,
        "The report's policy value lies in making heterogeneity legible without turning it into a simplistic ranking. For LERU and INTE, the practical opportunity is to validate the route map, agree a minimal vocabulary for public accountability and create a durable process for keeping institutional information current. That would convert a one-off mapping exercise into a shared European learning infrastructure.",
    )

    add_heading(doc, "Appendices and editorial work plan", level=1, page_break=True)
    add_heading(doc, "Recommended appendices", level=2)
    appendices = [
        "Appendix A: Concepts, variables and coding rules.",
        "Appendix B: Country system typology and 49-country status table.",
        "Appendix C: Transparency coding guide and 35-jurisdiction table.",
        "Appendix D: LERU 24-member route and output matrix.",
        "Appendix E: Member-validation questionnaire and change log.",
        "Appendix F: Source-selection protocol and limitations register.",
    ]
    for item in appendices:
        add_bullet(doc, bullet_num_id, item)
    add_heading(doc, "Editorial production plan", level=2)
    add_editorial_table(doc)

    add_heading(doc, "References and evidence plan", level=2)
    add_body(
        doc,
        "The final report should use a formal reference style and cite primary official sources at claim level. The current outline is grounded in the project's structured evidence backbone; the next drafting pass should convert the source registry and profile links into footnotes or endnotes, then freeze a bibliographic export at the publication cut-off.",
    )
    source_items = [
        ("LERU members page", "https://www.leru.org/members"),
        ("Project country dataset", "data/countries.js"),
        ("Project transparency dataset", "data/transparency.js"),
        ("Project source registry", "data/source-registry.csv"),
        ("LERU member dataset", "data/leru-members.js"),
        ("LERU public-source report", "reports/leru-institutions-report.html"),
        ("LERU institution extraction workflow", "docs/workflows/LERU-INSTITUTION-EXTRACTION-WORKFLOW.md"),
        ("Committee and case-file workflow", "docs/workflows/CASE-FILE-WORKFLOW.md"),
        ("Quality-consolidation plan", "docs/plans/NEXT-PHASE-QUALITY-PLAN.md"),
    ]
    for label, target in source_items:
        paragraph = doc.add_paragraph(style="Normal")
        apply_num(paragraph, bullet_num_id)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if target.startswith("http"):
            add_hyperlink(paragraph, label, target)
        else:
            lead = paragraph.add_run(label + ": ")
            set_run_font(lead, size=11, color=NAVY, bold=True)
            path_run = paragraph.add_run(target)
            set_run_font(path_run, name="Consolas", size=9.5, color=INK)

    add_callout(
        doc,
        "Review instruction.",
        "Use Word comments for decisions and Track Changes for textual revisions. Do not accept revisions or remove comments until the report format, authorship, data cut-off and member-validation route have been agreed.",
    )

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Build the LERU report/article outline DOCX.")
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()
    build_document(args.out.resolve())
    print(args.out.resolve())


if __name__ == "__main__":
    main()
