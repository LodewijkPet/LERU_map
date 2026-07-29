from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# compact_reference_guide preset
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "222222"
MUTED = "5B6570"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F4F6F9"
MID_GRAY = "D8DEE6"
WHITE = "FFFFFF"
CODE_BORDER = "D5DBE3"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(
                cell,
                top=CELL_TOP_BOTTOM_DXA,
                start=CELL_START_END_DXA,
                bottom=CELL_TOP_BOTTOM_DXA,
                end=CELL_START_END_DXA,
            )


def set_table_borders(table, color=MID_GRAY, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color=BLUE, size="18", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def set_run_font(
    run,
    *,
    name="Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*))"
)


def add_inline(paragraph, text: str, *, base_size: float | None = None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=base_size)

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(base_size or 10.5) - 0.5, color=DARK_BLUE)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                label, target = link_match.groups()
                if target.startswith(("http://", "https://")):
                    add_hyperlink(paragraph, label, target)
                else:
                    run = paragraph.add_run(label)
                    set_run_font(run, size=base_size, color=DARK_BLUE, bold=True)
                    path_run = paragraph.add_run(f" ({target})")
                    set_run_font(path_run, name="Consolas", size=(base_size or 10.5) - 1, color=MUTED)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, italic=True)
        position = match.end()

    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=base_size)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    if "Code Block" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Code Block"]
    style.font.name = "Consolas"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    style.font.size = Pt(8.5)
    style.font.color.rgb = rgb(INK)
    style.paragraph_format.left_indent = Inches(0.16)
    style.paragraph_format.right_indent = Inches(0.08)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_together = True

    if "Document Quote" not in [style.name for style in doc.styles]:
        quote = doc.styles.add_style("Document Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = doc.styles["Document Quote"]
    quote.font.name = "Calibri"
    quote._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = rgb(DARK_BLUE)
    quote.font.italic = True
    quote.paragraph_format.left_indent = Inches(0.28)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.2


def configure_section(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run(short_title)
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run("LERU Integrity Map | Agent Planning")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    label = footer_paragraph.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    add_field(footer_paragraph, "PAGE")


def add_cover(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    date_label: str,
    document_label: str,
) -> None:
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("LERU INTEGRITY MAP | AGENT PLANNING")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_paragraph.paragraph_format.keep_together = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, size=28, color=NAVY, bold=True)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(34)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_run_font(subtitle_run, size=14, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    meta_run = meta.add_run(date_label)
    set_run_font(meta_run, size=10.5, color=MUTED, bold=True)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(0)
    descriptor_run = descriptor.add_run(document_label)
    set_run_font(descriptor_run, size=9.5, color=MUTED, italic=True)

    doc.add_page_break()


def next_identifier(element, tag_name: str, attribute: str) -> int:
    values: list[int] = []
    for node in element.findall(qn(tag_name)):
        value = node.get(qn(attribute))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def create_abstract_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_identifier(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if bullet else f"%{level + 1}.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left" if bullet else "right")
        lvl.extend([start, num_fmt, lvl_text, lvl_jc])

        marker = 269 + (level * 360)
        text_indent = 540 + (level * 360)
        hanging = 271
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(text_indent))
        indent.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, indent, spacing])
        lvl.append(p_pr)

        if bullet:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Calibri")
            fonts.set(qn("w:hAnsi"), "Calibri")
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    return abstract_id


def create_number_instance(
    doc: Document,
    abstract_id: int,
    *,
    start_at: int | None = None,
    level: int = 0,
) -> int:
    numbering = doc.part.numbering_part.element
    num_id = next_identifier(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    if start_at is not None:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(max(0, min(level, 3))))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(max(0, start_at)))
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 3))))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def parse_table_row(line: str) -> list[str]:
    content = line.strip()
    if content.startswith("|"):
        content = content[1:]
    if content.endswith("|"):
        content = content[:-1]
    return [cell.strip() for cell in content.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def choose_column_widths(rows: list[list[str]]) -> list[int]:
    column_count = max(len(row) for row in rows)
    lengths = []
    for idx in range(column_count):
        max_length = max((len(re.sub(r"[`*_]", "", row[idx])) if idx < len(row) else 0) for row in rows)
        lengths.append(max(8.0, min(45.0, max_length) ** 0.72))

    if column_count == 2 and lengths[1] > lengths[0] * 1.35:
        return [2700, CONTENT_WIDTH_DXA - 2700]

    minimum = 1050 if column_count >= 4 else 1300
    available = CONTENT_WIDTH_DXA - (minimum * column_count)
    if available < 0:
        base = CONTENT_WIDTH_DXA // column_count
        widths = [base] * column_count
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
        return widths

    total_weight = sum(lengths)
    widths = [minimum + int(available * (weight / total_weight)) for weight in lengths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    widths = choose_column_widths(normalized)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])

    font_size = 9.0 if column_count <= 3 else 8.25
    for row_idx, row in enumerate(normalized):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade_cell(cell, PALE_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.paragraph_format.widow_control = True
            add_inline(paragraph, text, base_size=font_size)
            for run in paragraph.runs:
                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines or [""]):
        paragraph = doc.add_paragraph(style="Code Block")
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(5)
        if index == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(7)
        set_paragraph_shading(paragraph, PALE_GRAY)
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=8.5, color=INK)


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Document Quote")
    set_paragraph_left_border(paragraph)
    add_inline(paragraph, text, base_size=10.5)


def clean_heading(text: str) -> str:
    return re.sub(r"\s+#+\s*$", "", text.strip())


def add_body_from_markdown(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    start = next((idx for idx, line in enumerate(lines) if line.startswith("## ")), 0)
    lines = lines[start:]

    bullet_abstract = create_abstract_numbering(doc, bullet=True)
    decimal_abstract = create_abstract_numbering(doc, bullet=False)
    current_list_type: str | None = None
    current_num_id: int | None = None

    def reset_list() -> None:
        nonlocal current_list_type, current_num_id
        current_list_type = None
        current_num_id = None

    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        if text:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, text)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            reset_list()
            index += 1
            continue

        fence_match = re.match(r"^\s*(`{3,})(.*)$", line)
        if fence_match:
            flush_paragraph()
            reset_list()
            fence_length = len(fence_match.group(1))
            code_lines: list[str] = []
            index += 1
            while index < len(lines):
                close = re.match(r"^\s*(`{3,})\s*$", lines[index])
                if close and len(close.group(1)) >= fence_length:
                    break
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            reset_list()
            table_rows = [parse_table_row(lines[index])]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[index]))
                index += 1
            add_markdown_table(doc, table_rows)
            continue

        heading_match = re.match(r"^(#{2,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            reset_list()
            source_level = len(heading_match.group(1))
            word_level = max(1, min(3, source_level - 1))
            paragraph = doc.add_paragraph(style=f"Heading {word_level}")
            add_inline(paragraph, clean_heading(heading_match.group(2)))
            index += 1
            continue

        quote_match = re.match(r"^\s*>\s?(.*)$", line)
        if quote_match:
            flush_paragraph()
            reset_list()
            quote_lines = [quote_match.group(1)]
            index += 1
            while index < len(lines):
                next_quote = re.match(r"^\s*>\s?(.*)$", lines[index])
                if not next_quote:
                    break
                quote_lines.append(next_quote.group(1))
                index += 1
            add_quote(doc, " ".join(quote_lines))
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.+)$", line)
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if bullet_match or number_match:
            flush_paragraph()
            list_type = "bullet" if bullet_match else "number"
            match = bullet_match or number_match
            level = min(len(match.group(1).replace("\t", "    ")) // 2, 3)
            if list_type == "number":
                current_num_id = create_number_instance(
                    doc,
                    decimal_abstract,
                    start_at=int(number_match.group(2)),
                    level=level,
                )
                current_list_type = "number"
                item_text = number_match.group(3)
            else:
                item_text = bullet_match.group(2)
            if list_type == "bullet" and (current_list_type != list_type or current_num_id is None):
                abstract = bullet_abstract if list_type == "bullet" else decimal_abstract
                current_num_id = create_number_instance(doc, abstract)
                current_list_type = list_type
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, current_num_id, level)
            add_inline(paragraph, item_text)
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def audit_document(doc: Document) -> list[str]:
    findings: list[str] = []
    section = doc.sections[0]
    expected_margin = int(Inches(MARGIN_IN))
    for name, value in (
        ("top", section.top_margin),
        ("right", section.right_margin),
        ("bottom", section.bottom_margin),
        ("left", section.left_margin),
    ):
        if int(value) != expected_margin:
            findings.append(f"{name} margin does not match preset")

    normal = doc.styles["Normal"]
    if normal.font.name != "Calibri":
        findings.append("Normal font is not Calibri")
    if normal.font.size != Pt(11):
        findings.append("Normal size is not 11 pt")

    for table_index, table in enumerate(doc.tables, start=1):
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid]
        if sum(widths) != CONTENT_WIDTH_DXA:
            findings.append(f"Table {table_index} grid width is {sum(widths)}, expected {CONTENT_WIDTH_DXA}")
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None or int(tc_w.get(qn("w:w"))) != widths[idx]:
                    findings.append(f"Table {table_index} has inconsistent cell geometry")
                    break

    return sorted(set(findings))


def build_document(
    source_path: Path,
    output_path: Path,
    *,
    title: str,
    subtitle: str,
    date_label: str,
    short_title: str,
    document_label: str,
) -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc, short_title)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "LERU Integrity Map project team"
    doc.core_properties.keywords = "agent plans; research integrity; workflow; documentation"
    doc.core_properties.comments = "English-language reference document generated from the project Markdown source."

    add_cover(
        doc,
        title=title,
        subtitle=subtitle,
        date_label=date_label,
        document_label=document_label,
    )
    add_body_from_markdown(doc, source_path.read_text(encoding="utf-8"))

    findings = audit_document(doc)
    if findings:
        raise RuntimeError("Preset audit failed:\n- " + "\n- ".join(findings))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build polished DOCX files from Agent-planning Markdown.")
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument("--subtitle", required=True)
    parser.add_argument("--date-label", required=True)
    parser.add_argument("--short-title", required=True)
    parser.add_argument("--document-label", required=True)
    args = parser.parse_args()

    build_document(
        args.source.resolve(),
        args.output.resolve(),
        title=args.title,
        subtitle=args.subtitle,
        date_label=args.date_label,
        short_title=args.short_title,
        document_label=args.document_label,
    )
    print(args.output.resolve())


if __name__ == "__main__":
    main()
