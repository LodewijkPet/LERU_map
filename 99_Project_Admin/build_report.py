from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "LERU_Research_Integrity_Committees_Comparative_Report_WORKING_DRAFT.docx"
COVERAGE_PATH = ROOT / "00_Shared_LERU_and_Methods" / "04_Cross_Committee_Data" / "CURRENT_DOSSIER_COVERAGE.json"

NAVY = "203748"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
GOLD = "C28D2C"
MUTED = "657785"
LIGHT = "F4F6F9"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "F8F2E6"
BORDER = "B7C1CC"
WHITE = "FFFFFF"
BLACK = "202020"

BODY_FONT = "Calibri"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


COUNTRIES = [
    ("Belgium", "BE", ["KU Leuven"]),
    ("Denmark", "DK", ["University of Copenhagen"]),
    ("Finland", "FI", ["University of Helsinki"]),
    ("France", "FR", ["Paris-Saclay University", "Sorbonne University", "University of Strasbourg"]),
    ("Germany", "DE", ["Heidelberg University", "LMU Munich", "University of Freiburg"]),
    ("Ireland", "IE", ["Trinity College Dublin"]),
    ("Italy", "IT", ["University of Milan"]),
    ("Netherlands", "NL", ["Leiden University", "University of Amsterdam", "Utrecht University"]),
    ("Spain", "ES", ["University of Barcelona"]),
    ("Sweden", "SE", ["Lund University"]),
    ("Switzerland", "CH", ["ETH Zurich", "University of Geneva", "University of Zurich"]),
    ("United Kingdom", "UK", ["Imperial College London", "University College London", "University of Cambridge", "University of Edinburgh", "University of Oxford"]),
]

COUNTRY_SYSTEMS = {
    "Belgium": (
        "Working classification: a federal and linguistically differentiated system in which research-integrity norms and support are distributed across national, Flemish and French-speaking structures. Institutional procedures remain the principal route for individual cases."
    ),
    "Denmark": (
        "Working classification: a statutory two-track system. The national level addresses research misconduct within its legal remit, while universities address questionable research practices and maintain local advisory and screening structures."
    ),
    "Finland": (
        "Working classification: national self-regulation coordinated by the Finnish National Board on Research Integrity (TENK). Committed institutions conduct the first inquiry under the national guideline; TENK can review procedural compliance and issue opinions."
    ),
    "France": (
        "Working classification: a legally codified framework implemented through institutional scientific-integrity officers (référents à l’intégrité scientifique, RIS), with national coordination and guidance by the French Office for Research Integrity (Ofis)."
    ),
    "Germany": (
        "Working classification: federal and institutionally self-regulated, shaped by the German Research Foundation (DFG) Code, ombudsperson systems and institutional investigation commissions; sectoral and funder routes may operate alongside university procedures."
    ),
    "Ireland": (
        "Working classification: harmonised soft law and sector coordination, implemented by research-performing organisations through named Research Integrity Officers and institutional procedures, with national forums supporting common standards."
    ),
    "Italy": (
        "Working classification: a distributed system in which institutional ethics, integrity and disciplinary bodies interact with national law, funders and para-national actors; a single universal national misconduct route has not yet been confirmed."
    ),
    "Netherlands": (
        "Working classification: soft-law regulation through the Netherlands Code of Conduct for Research Integrity. Institutions investigate and decide in the first instance; the Netherlands Board on Research Integrity (LOWI) provides a national second-opinion route before final institutional decisions."
    ),
    "Spain": (
        "Working classification: a multi-layered and partly consultative landscape in which national, regional and institutional ethics/integrity bodies may have overlapping advisory, preventive and case-handling roles."
    ),
    "Sweden": (
        "Working classification: a statutory two-track model. The National Board for Assessment of Research Misconduct (NPOF) examines legally defined fabrication, falsification and plagiarism; universities examine other deviations from good research practice."
    ),
    "Switzerland": (
        "Working classification: a federal and institutionally autonomous system shaped by a national code, national coordination bodies and funder requirements. Universities retain their own ombudsperson, delegate, commission and executive decision routes."
    ),
    "United Kingdom": (
        "Working classification: employer-led implementation under the Concordat to Support Research Integrity and funder-assurance requirements. Universities maintain local reporting, screening, investigation and annual-statement arrangements, supported by sector guidance such as UKRIO and Russell Group cooperation."
    ),
}

PROFILES = {
    "KU Leuven": {
        "route": "Commission on Research Integrity (CRI), Research Integrity Reporting Desk and research-integrity advisers",
        "position": "The Reporting Desk is embedded in the Research Integrity and Ethics Unit of the Research Coordination Office; the CRI handles formal complaints.",
        "decision": "The CRI investigates and issues reasoned findings/recommendations; follow-up outside the CRI procedure sits with the competent university authorities.",
        "interfaces": "Research Policy/Coordination Office; group advisers; confidential support; ethics committees; funders and external authorities where applicable.",
        "gaps": "Confirm current composition, appointment terms, complete stage-specific timing, five-year caseload/outcomes and the exact publication series."
    },
    "University of Copenhagen": {
        "route": "Committee for Responsible Research Practice and related local advisers/named persons",
        "position": "Rector-appointed university committee; secretariat support is provided within University Administration by Research and Information Security.",
        "decision": "The local committee addresses questionable research practice and routes matters falling within the statutory misconduct remit to the national board.",
        "interfaces": "Rector; University Administration; Danish Committee on Research Misconduct; local faculties and advisers.",
        "gaps": "Verify committee composition and term history, current division of triage responsibilities, local outcome categories and annual case counts."
    },
    "University of Helsinki": {
        "route": "Chancellor-led inquiry route, supported by research-integrity advisers",
        "position": "The Chancellor is the university-level authority for alleged violations; the administrative placement of advisers and case support is not explicit in the current public baseline.",
        "decision": "The institution conducts the TENK-aligned inquiry and reports as required; TENK can review whether the national procedure was followed.",
        "interfaces": "Chancellor; research-integrity advisers; legal/academic affairs; TENK; faculties and research units.",
        "gaps": "Confirm the office that coordinates files, inquiry-panel composition, published case statistics, current adviser network and document-version history."
    },
    "Paris-Saclay University": {
        "route": "POLETHIS research ethics/scientific-integrity council and network of scientific-integrity referents",
        "position": "Preliminary mapping places the function with or close to the University Presidency; referents receive reports and promote integrity.",
        "decision": "Referents investigate reports within their remit and advise university leadership; the precise allocation of final decisions and disciplinary follow-up requires confirmation.",
        "interfaces": "University Presidency; POLETHIS; component institutions/laboratories; Ofis; doctoral schools and ethics structures.",
        "gaps": "Map the multi-component university route, establish which procedure governs each population, and obtain composition, timing, output and caseload data."
    },
    "Sorbonne University": {
        "route": "Scientific Integrity Delegation led by the RIS, with a Scientific Integrity Committee and ambassadors",
        "position": "The RIS is appointed by the President and reports to academic/governance bodies; a single administrative host unit is not yet explicit.",
        "decision": "The RIS investigates and prepares reports/recommendations; institutional leadership decides on scientific, disciplinary or organisational follow-up.",
        "interfaces": "President; Academic Council/Board; Scientific Integrity Committee; ambassador network; faculties; Ofis.",
        "gaps": "Confirm institution-wide versus faculty-specific procedures, administrative support, committee composition, annual-report series and comparable outcome counts."
    },
    "University of Strasbourg": {
        "route": "Scientific Integrity Referent",
        "position": "The referent receives and investigates reports and submits recommendations to the University President; no separate administrative home is confirmed.",
        "decision": "The referent conducts an adversarial inquiry and writes a report; the President decides on measures.",
        "interfaces": "University President; external experts; Ofis; LERU integrity network; legal/disciplinary and ethics routes where relevant.",
        "gaps": "Obtain the formal procedure and mandate, appointment/independence safeguards, timing rules, annual activity and document history."
    },
    "Heidelberg University": {
        "route": "Senate Commission for Safeguarding Good Academic Practice and ombudspersons",
        "position": "University governance under Senate/Rectorate; Legal Affairs appears as an administrative contact, but a single host office is not explicit.",
        "decision": "Ombudspersons handle advice/preliminary matters; a standing commission investigates suspected misconduct under the university rules.",
        "interfaces": "Senate/Rectorate; Legal Affairs; ombudspersons; DFG/central ombuds structures; faculties and disciplinary authorities.",
        "gaps": "Validate current membership, administrative support, decision/sanction route, case counts, publication practice and changes since the 2021 rules."
    },
    "LMU Munich": {
        "route": "Ombudspersons, officer for self-control in science and Investigation Committee",
        "position": "Anchored in university governance/Senate arrangements; no single operational office is stated in the preliminary public evidence.",
        "decision": "The ombuds/self-control route performs preliminary handling; the Investigation Committee conducts formal examination and leadership determines measures.",
        "interfaces": "Senate; university leadership; ombudspersons; Investigation Committee; DFG/sector bodies; faculties.",
        "gaps": "Reconcile the 2023/2025 governance instruments, confirm current appointments and secretariat, and obtain activity/transparency data."
    },
    "University of Freiburg": {
        "route": "Coordination Office for Research Integrity, ombudsperson/self-control route and investigation commission",
        "position": "The Coordination Office was established by the Rectorate and is supported by an adviser to the Pro-Rector responsible for research integrity.",
        "decision": "The coordination and ombuds route supports intake/prevention; formal bodies investigate and university leadership follows up.",
        "interfaces": "Rectorate/Pro-Rector; Coordination Office; ombudspersons; investigation commission; DFG and external ombuds structures.",
        "gaps": "Confirm establishment timeline, role separation, composition, procedural time limits, public outputs and yearly allegations/outcomes."
    },
    "Trinity College Dublin": {
        "route": "Office of the Dean of Research, Associate Deans and named Research Integrity Officer/Senior Dean route",
        "position": "Initial coordination sits with the Office of the Dean of Research; disciplinary and population-specific routing involves senior academic officers.",
        "decision": "Preliminary enquiries are coordinated by the Dean of Research; serious/formal matters move through the applicable college investigation and disciplinary structures.",
        "interfaces": "Dean of Research; Senior Dean/RIO; College Secretary; faculty Associate Deans; national Research Integrity Forum and funders.",
        "gaps": "Confirm the current end-to-end route, role titles, appeal route, annual statistics and how student, staff and public reports are separated."
    },
    "University of Milan": {
        "route": "Research Ethics Office, Ethics Committee and separate disciplinary bodies",
        "position": "The ethics support function appears within a staff unit for privacy and research ethics; disciplinary proceedings sit in Legal Affairs, with governance-level committee roles.",
        "decision": "Responsibilities appear split between advice/ethics review and formal employment-disciplinary handling; the primary research-misconduct investigation route remains to be validated.",
        "interfaces": "Vice-Rector/research governance; Ethics Committee; Research Ethics Office; Disciplinary Board/UPD; Legal Affairs; national/funder actors.",
        "gaps": "Highest-priority validation: identify the primary misconduct route, its mandate, procedure, composition, public outputs and relation to the Ethics Committee."
    },
    "Leiden University": {
        "route": "Academic Integrity Committee (CWI), with Leiden University and LUMC chambers, plus confidential advisers",
        "position": "University-wide committee advising the Executive Board; secretariat support is located in Legal Affairs.",
        "decision": "The CWI investigates and advises; the Executive Board issues an initial decision, with possible LOWI review before the final decision.",
        "interfaces": "Executive Board; Legal Affairs; confidential advisers; LUMC; LOWI; Universities of the Netherlands publication route.",
        "gaps": "Complete the committee/secretariat history, current composition, annual counts by complaint type/outcome, durations and cross-institutional arrangements."
    },
    "University of Amsterdam": {
        "route": "Scientific/Academic Integrity Committee (CWI/SIC) and confidential adviser",
        "position": "University-wide committee appointed by the Executive Board and administratively supported by Legal Affairs.",
        "decision": "The committee investigates and advises; the Executive Board decides, subject to possible LOWI advice before finalisation.",
        "interfaces": "Executive Board; Legal Affairs; confidential adviser; LOWI; Universities of the Netherlands; collaborating institutions.",
        "gaps": "Update the regulation/version history, extract the 2019–2024 caseload consistently, verify composition terms and code cross-institutional cases."
    },
    "Utrecht University": {
        "route": "Committee for Research Integrity and confidential advisers",
        "position": "Established by the Executive Board and supported by Legal Affairs.",
        "decision": "The committee investigates and advises; the Executive Board decides, with LOWI as the national second-opinion route.",
        "interfaces": "Executive Board; Legal Affairs; confidential advisers; LOWI; Universities of the Netherlands; partner institutions.",
        "gaps": "Obtain and normalise annual statistics, current membership, procedure dates, publication coverage and collaboration/referral rules."
    },
    "University of Barcelona": {
        "route": "University research-integrity route involving the Vice-Rector for Research and an integrity committee; the Bioethics Commission provides related policy support",
        "position": "The preliminary record does not identify a single administrative office; reports appear to enter through research leadership and committee structures report to the Rector.",
        "decision": "The exact division between integrity assessment, ethics advice and executive/disciplinary decision-making is not yet verified.",
        "interfaces": "Vice-Rector/Rector; institutional integrity and bioethics bodies; Catalan CIR-CAT; faculties; disciplinary structures.",
        "gaps": "Distinguish misconduct handling from ethics review, obtain the formal complaint procedure, and confirm membership, authority, outputs and statistics."
    },
    "Lund University": {
        "route": "Deviations from Good Research Practice Review Board",
        "position": "University-management route under the Vice-Chancellor; the Board examines other deviations and supports decisions within the university.",
        "decision": "Legally defined FFP is examined nationally by NPOF; other deviations are handled by Lund's Board and decided/followed up within the university.",
        "interfaces": "Vice-Chancellor; Review Board; faculty research representatives; NPOF; departments and national authorities.",
        "gaps": "Extract the 2025 annual report, composition and terms, phase-specific durations, referral patterns and outcome categories into the common dataset."
    },
    "ETH Zurich": {
        "route": "Scientific Integrity Office supporting the Integrity Commission, with advisers and good-scientific-practice structures",
        "position": "Located in the Executive Board domain of Research; the Scientific Integrity Office supports case handling and the Commission conducts investigations.",
        "decision": "The Commission investigates and reports within the ETH governance framework; executive/disciplinary authorities determine applicable follow-up.",
        "interfaces": "Vice President/Executive Board for Research; Integrity Commission; Scientific Integrity Office; advisers; Swiss code/funders; legal and disciplinary functions.",
        "gaps": "Integrate the received ETH documents, verify permission to cite email evidence, and extract composition, annual counts, outcomes and document evolution."
    },
    "University of Geneva": {
        "route": "Rectorate-led integrity procedure with faculty delegates and fact-finding/investigation arrangements",
        "position": "Integrity procedures are centred on the Rectorate; Legal Affairs attached to the Rectorate is listed as responsible for integrity directives.",
        "decision": "Delegates/intake actors support assessment and investigation; the Rectorate determines institutional follow-up under the directive.",
        "interfaces": "Rectorate; Legal Affairs; faculty delegates; ad-hoc experts/commission; Swiss national code and sector bodies.",
        "gaps": "Confirm current directive version, composition/appointment of investigation bodies, appeal route, public reporting and yearly case activity."
    },
    "University of Zurich": {
        "route": "Research Integrity Coordination Office, ombudspersons and research-integrity delegates",
        "position": "The Coordination Office is within the Office of the Vice President Research/Research and Grants Office—the placement that triggered the original comparative request.",
        "decision": "The Coordination Office initiates and supports the process; ombudsperson/delegate and executive routes perform defined assessment, investigation and decision roles.",
        "interfaces": "Vice President Research; Research and Grants Office; ombudspersons; delegates; Executive Board; Swiss codes/funders and legal functions.",
        "gaps": "Validate the desired/current organisational placement, separation from Legal Affairs, staffing, case statistics, decision routes and recent ordinance changes."
    },
    "Imperial College London": {
        "route": "Research Misconduct Review Group and Research Governance/Integrity support, with a Director of Research Integrity Investigations",
        "position": "Integrity support sits in the Research Office; the formal review group's organisational home is less explicit than the operational case contact.",
        "decision": "Regulation 21 provides screening, investigation and institutional decision/follow-up arrangements; annual statements report high-level activity.",
        "interfaces": "Research Office; senior governance; HR/legal; funders; collaborating institutions; UKRIO/Russell Group guidance.",
        "gaps": "Extract the February 2026 revision, reconcile historical annual statements, confirm membership/appointment and code comparable durations/outcomes."
    },
    "University College London": {
        "route": "Research Misconduct Committee and Governance Team, with separate research-integrity support and student-case teams",
        "position": "Functions are distributed across Research & Innovation Services, the Office of General Counsel/Governance, and Student and Registry Services.",
        "decision": "The misconduct procedure uses screening and investigation panels, with institutional governance responsible for decisions and follow-up; student academic misconduct is separate.",
        "interfaces": "Research Misconduct Committee; Registrar/governance; Research & Innovation Services; student casework; third-party notification; funders/partners.",
        "gaps": "Obtain the current SharePoint-hosted procedure files, verify role allocation, extract annual statements, composition and cross-institutional workflow."
    },
    "University of Cambridge": {
        "route": "Research Governance and Integrity Team and misconduct investigation/advisory panels",
        "position": "The support team sits in the University/Cambridge Research Office; student conduct is handled separately within education services.",
        "decision": "The published procedure defines preliminary assessment and formal investigation, routed through senior research governance and university authorities.",
        "interfaces": "Research Office; Pro-Vice-Chancellor for Research; panel members; HR/legal; OSCCA for student boundaries; funders and collaborators.",
        "gaps": "Extract current panel rules, exact reporting lines, annual statement counts, duration data and the relationship between advisory and investigative bodies."
    },
    "University of Edinburgh": {
        "route": "Research Governance, Compliance and Risk team with College Named Persons and separate student misconduct bodies",
        "position": "Research-integrity support sits in the Edinburgh Research Office; case initiation and handling involve college-level Named Persons and central governance.",
        "decision": "The institutional procedure routes allegations through named and senior decision-makers; student misconduct follows separate academic processes.",
        "interfaces": "Edinburgh Research Office; colleges/Named Persons; senior university officers; HR/legal; student misconduct bodies; funders and partners.",
        "gaps": "Verify current procedure and Named Person network, extract 2025–26 statement data, composition/appointment, timelines and appeal/funder notifications."
    },
    "University of Oxford": {
        "route": "Registrar-led research-integrity procedure supported by Research Services and investigation/advisory panels",
        "position": "Research Services provides policy/support; the Registrar is the named senior officer for formal allegations. Student academic misconduct is separate under the Proctors.",
        "decision": "The current procedure provides proportional assessment, formal panels and an Appeals Panel, with institutional leaders responsible for decisions and measures.",
        "interfaces": "Registrar; Research Services; departments/faculties; panels; Pro-Vice-Chancellor; Proctors boundary; funders, regulators and partners.",
        "gaps": "Archive the current 2026 web procedure/SSO guides, confirm commencement/version, collect annual statement data and extract panel composition/time limits."
    },
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, name=BODY_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def configure_table(table, widths_dxa, *, indent=TABLE_INDENT_DXA, header=True, font_size=9.1):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        if header and row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[col_idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[col_idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, color=BLACK)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size, color=NAVY, bold=True)
    set_table_borders(table)


def add_table(doc, headers, rows, widths_dxa, *, font_size=9.1, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
        set_cell_shading(table.rows[0].cells[idx], header_fill)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = "" if value is None else str(value)
    configure_table(table, widths_dxa, header=True, font_size=font_size)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
    return table


def add_key_value_table(doc, rows, *, label_width=2700, font_size=9.4):
    value_width = CONTENT_WIDTH_DXA - label_width
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT)
    configure_table(table, [label_width, value_width], header=False, font_size=font_size)
    for row in table.rows:
        set_cell_shading(row.cells[0], LIGHT)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=font_size, color=NAVY, bold=True)
    return table


def add_callout(doc, title, text, *, fill=PALE_BLUE, title_color=DEEP_BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.2, color=BLACK)
    configure_table(table, [CONTENT_WIDTH_DXA], header=False, font_size=10.2)
    set_cell_shading(cell, fill)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=9.2, color=MUTED, italic=True)
    return p


def add_source_note(doc, text):
    p = doc.add_paragraph(style="Table Source")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=8.7, color=MUTED, italic=True)
    return p


def add_body(doc, text, *, bold_lead=None, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=BLACK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=BLACK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK, italic=italic)
    return p


def add_list_item(doc, text, *, numbered=False, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=BLACK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    return p


def add_page_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field if this contents list is not populated."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def paragraph_bottom_border(paragraph, color=GOLD, size="12", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_section_page_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_running_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header.distance = Inches(0.492)
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("LERU Research Integrity Functions")
    set_run_font(r, size=8.6, color=MUTED, bold=True)
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    r2 = hp.add_run("\tWORKING DRAFT · 18 AUGUST 2026")
    set_run_font(r2, size=8.4, color=MUTED)

    footer = section.footer
    footer.distance = Inches(0.492)
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r3 = fp.add_run("Project-team working document  ·  Page ")
    set_run_font(r3, size=8.3, color=MUTED)
    page_run = add_page_field(fp, "PAGE")
    set_run_font(page_run, size=8.3, color=MUTED)
    r4 = fp.add_run(" of ")
    set_run_font(r4, size=8.3, color=MUTED)
    total_run = add_page_field(fp, "NUMPAGES")
    set_run_font(total_run, size=8.3, color=MUTED)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1

    def add_definition(abstract_id, num_id, fmt, text_value, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(9):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), fmt)
            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text_value if fmt == "bullet" else f"%{level + 1}.")
            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "279")
            p_pr.extend([tabs, ind])
            lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
            if font:
                r_pr = OxmlElement("w:rPr")
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), font)
                r_fonts.set(qn("w:hAnsi"), font)
                r_pr.append(r_fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", BODY_FONT)
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    caption.font.size = Pt(9.2)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(BLACK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Table Source" not in styles:
        source_style = styles.add_style("Table Source", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source_style = styles["Table Source"]
    source_style.font.name = BODY_FONT
    source_style.font.size = Pt(8.7)
    source_style.font.italic = True
    source_style.font.color.rgb = rgb(MUTED)
    source_style.paragraph_format.space_before = Pt(4)
    source_style.paragraph_format.space_after = Pt(4)

    toc_styles = ["TOC 1", "TOC 2", "TOC 3"]
    for idx, name in enumerate(toc_styles, start=1):
        if name in styles:
            style = styles[name]
            style.font.name = BODY_FONT
            style.font.size = Pt(10.2 if idx == 1 else 9.6)
            style.font.color.rgb = rgb(NAVY if idx == 1 else BLACK)
            style.paragraph_format.space_after = Pt(2)


def read_coverage():
    data = json.loads(COVERAGE_PATH.read_text(encoding="utf-8-sig"))
    return {item["institution"]: item for item in data["dossiers"]}


def add_profile(doc, institution, coverage):
    profile = PROFILES[institution]
    cov = coverage.get(institution, {})
    add_heading(doc, institution, 3)
    evidence = (
        f"Local dossier: {cov.get('local_file_count', '—')} indexed files; "
        f"{cov.get('procedure_or_policy_files', '—')} procedure/policy file(s); "
        f"{cov.get('report_statistics_timeline_files', '—')} report/statistics/timeline file(s); "
        f"{cov.get('restricted_correspondence_files', '—')} restricted correspondence file(s). "
        "File presence does not establish substantive completeness or currency."
    )
    add_key_value_table(doc, [
        ("Primary route/function", profile["route"]),
        ("Institutional placement", profile["position"]),
        ("Investigative/decision logic", profile["decision"]),
        ("Direct interfaces", profile["interfaces"]),
        ("Evidence archived", evidence),
        ("Priority validation", profile["gaps"]),
        ("Current interpretation", "Preliminary working synthesis. Not yet confirmed by the institution and not suitable for ranking."),
    ], label_width=2450, font_size=9.1)
    add_source_note(doc, "Source basis: official public material, legacy structured searches and, where available, locally preserved correspondence. Each statement must be re-coded with field-level provenance before the comparative results are finalised.")


def add_reference(doc, number, title, url=None, note=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{number}. ")
    set_run_font(r, size=9.5, color=BLACK, bold=True)
    r2 = p.add_run(title)
    set_run_font(r2, size=9.5, color=BLACK)
    if url:
        p.add_run(" ")
        add_hyperlink(p, url, url)
    if note:
        r3 = p.add_run(f" {note}")
        set_run_font(r3, size=9.5, color=MUTED, italic=True)


doc = Document()
doc.core_properties.title = "Research Integrity Committees and Functions across LERU Universities"
doc.core_properties.subject = "Comparative mapping protocol and working report"
doc.core_properties.author = "LERU Research Integrity Mapping project"
doc.core_properties.keywords = "LERU; research integrity; scientific integrity; committees; misconduct; comparative mapping"
doc.core_properties.comments = "Working draft generated 18 August 2026; requires member validation before external use."

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

configure_styles(doc)
coverage = read_coverage()

# Cover
cover_section = doc.sections[0]
cover_section.header.is_linked_to_previous = False
cover_section.footer.is_linked_to_previous = False
cover_section.header.paragraphs[0].text = ""
cover_section.footer.paragraphs[0].text = ""

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(104)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(18)
kr = kicker.add_run("LERU RESEARCH INTEGRITY MAPPING")
set_run_font(kr, size=10.5, color=GOLD, bold=True)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
tr = title.add_run("Research Integrity Committees\nand Functions across LERU Universities")
set_run_font(tr, size=30, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)
sr = subtitle.add_run("Comparative mapping protocol, evidence baseline and country-by-country working report")
set_run_font(sr, size=15, color=DEEP_BLUE)

rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(28)
paragraph_bottom_border(rule, color=GOLD, size="14", space="2")

status = doc.add_paragraph()
status.alignment = WD_ALIGN_PARAGRAPH.CENTER
status.paragraph_format.space_after = Pt(72)
st = status.add_run("WORKING DRAFT  ·  VERSION 0.1")
set_run_font(st, size=10.5, color=GOLD, bold=True)

datep = doc.add_paragraph()
datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
datep.paragraph_format.space_after = Pt(5)
dr = datep.add_run("18 August 2026")
set_run_font(dr, size=12, color=NAVY, bold=True)

prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
prep.paragraph_format.space_after = Pt(0)
pr = prep.add_run("Standalone report framework · 24 LERU member universities · 12 countries")
set_run_font(pr, size=9.5, color=MUTED, italic=True)

# Main section begins at page 1
main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
main_section.page_width = Inches(8.5)
main_section.page_height = Inches(11)
main_section.top_margin = Inches(1)
main_section.bottom_margin = Inches(1)
main_section.left_margin = Inches(1)
main_section.right_margin = Inches(1)
set_section_page_start(main_section, 1)
set_running_header_footer(main_section)

add_heading(doc, "Document control", 1)
add_key_value_table(doc, [
    ("Document", "Research Integrity Committees and Functions across LERU Universities"),
    ("Version and date", "Working draft 0.1 · 18 August 2026"),
    ("Population", "All 24 official LERU member universities in 12 countries"),
    ("Primary unit", "The institution's research-integrity handling route; nested route actors may include a committee, office, ombudsperson, referent, reporting desk, adviser or panel"),
    ("Current status", "Scientific report architecture and preliminary evidence baseline; comparative field extraction and member verification are not yet complete"),
    ("Confidentiality", "Public sources may be cited. Correspondence and meeting material remain restricted until permission and redaction have been checked."),
    ("Independence from website", "This Word report is the analytical master. Any later website is a dissemination product derived from a validated report dataset."),
], label_width=2500, font_size=9.4)

add_callout(
    doc,
    "Reader's note",
    "The current population is validated as 24 LERU member universities, not as 24 confirmed members of a single LERU-INTE committee network. For all institutions, INTE participation remains to be validated. The word ‘committee’ is therefore used as an accessible umbrella; the analysis itself follows the actual local route, which may be a committee, office, named officer, ombudsperson, referent or hybrid arrangement.",
    fill=PALE_GOLD,
    title_color=GOLD,
)

add_heading(doc, "Executive summary", 1)
add_body(doc, "This report addresses a concrete governance question: where is the university function that handles potential research misconduct, promotes research integrity and oversees relevant ordinances institutionally located, and how does that function operate? The original request concerned organisational placement. The present report expands that question into a systematic comparison of national context, institutional mandate, procedure, transparency, activity, related bodies and historical development across all LERU member universities.")
add_body(doc, "The project uses a descriptive comparative design. All 24 official LERU members are included; no sampling or performance ranking is intended. Public laws, codes, procedures, official webpages, annual reports and case outputs are combined with direct feedback received from institutional representatives. Large language models support multilingual searching, extraction and structuring, but are not treated as evidence. Every final claim must be traceable to an original source or an attributable member confirmation.")
add_body(doc, "The preliminary evidence base establishes an institutional route or procedure for all 24 universities. Fourteen profiles currently have strong and ten moderate public-source evidence. A strict indicator identifies public report or case-output evidence for 14 of 24 institutions. Transparency is heterogeneous: 13 institutions have institution-owned local output, four depend mainly on national or sector output, five currently expose procedure only, one has restricted/internal output and one has historical or case-specific output. These categories describe source visibility, not institutional quality.")
add_body(doc, "A complete dossier structure has now been created for every university, with separate areas for national context, institutional governance, procedures, statistics/timeline, restricted correspondence and web evidence. The key remaining work is normalised field-level extraction and a targeted validation round. The final report should therefore make only bounded preliminary observations until the same variables have been verified for every institution.")

add_callout(
    doc,
    "Primary research question",
    "How do the institutional placement, mandate, composition, case-handling procedures, transparency practices and system interfaces of research-integrity committees/functions differ across LERU member universities, and how are those differences shaped by national research-integrity systems?",
)

add_heading(doc, "Contents", 1)
add_toc(doc)
doc.add_page_break()

# 1 Introduction
add_heading(doc, "1. Introduction", 1)
add_heading(doc, "1.1 Background", 2)
add_body(doc, "Research-integrity concerns are handled at the interface of scientific norms, employment and disciplinary law, university governance, funder requirements, publication correction and national oversight. That interface is organised differently across Europe. In one university the visible front door may be a standing commission; elsewhere it is a rector, chancellor, named person, ombudsperson, scientific-integrity referent, research office or legal/governance team. A formally similar complaint may therefore enter a different organisational route, be screened against a different definition, and lead to a different sequence of advice, investigation, decision, review and publication.")
add_body(doc, "A meaningful comparison cannot stop at the name of a committee. It must reconstruct the complete route: the national normative system; the institutional position of the function; its mandate, composition and independence; the end-to-end procedure; the final decision and review route; the actors with which it exchanges information; and the documents and public outputs through which the system can be observed.")

add_heading(doc, "1.2 Why LERU is a useful comparative population", 2)
add_body(doc, "LERU is a bounded network of research-intensive universities operating across several European legal and governance traditions. As of the evidence cut-off, it comprises 24 member universities in 12 countries [1]. The population combines institutions subject to statutory national misconduct bodies, national self-regulatory systems, employer-led systems and institutionally autonomous models. It is therefore useful for examining how similar research-integrity responsibilities are embedded in different system architectures.")
add_body(doc, "LERU has also developed network-level work on research culture and integrity, including its 2020 paper on research-integrity culture and later work on responsible communication [3,4]. These materials make the network substantively relevant, while membership provides a transparent inclusion rule. The study does not assume that all members organise or label an ‘integrity committee’ in the same way.")

add_heading(doc, "1.3 The original request and the scope of this report", 2)
add_body(doc, "The original request, preserved in the project archive, asked where the office responsible for handling potential academic-misconduct cases, promoting academic integrity and overseeing the integrity ordinance is institutionally located at each university. The immediate practical context was the University of Zurich's review of the placement of its Research Integrity Coordination Office, which had emerged from and remained formally within the Research and Grants Office despite the legal character of many cases.")
add_body(doc, "The first response mapped institutional placement across 24 universities. The present report deliberately treats that work as a starting point rather than a completed result. Organisational placement is compared together with the national system, the actual handling route, procedure, composition, relationships, transparency, case activity and history. This makes the report independently readable and scientifically defensible before any connection is made to the existing website.")

add_heading(doc, "1.4 Objectives", 2)
for item in [
    "Describe the national research-integrity system in each of the 12 countries represented in LERU.",
    "Identify the primary institutional route for questions, reports and formal allegations at each LERU university.",
    "Determine where that route is organisationally situated and how its independence, secretariat and reporting line are arranged.",
    "Compare mandate, jurisdiction, composition, appointment, intake, admissibility, investigation, decision, measures, review, publication and protection of parties.",
    "Compile a versioned inventory of constitutive documents, procedures, annual reports, case outputs, related policies and received confirmations.",
    "Describe workload and outcomes where definitions and denominators permit responsible comparison.",
    "Reconstruct LERU-, country-, institution-, route- and document-level timelines.",
    "Identify common models, meaningful differences, data gaps and practical implications without ranking institutions.",
]:
    add_list_item(doc, item, numbered=True)

add_heading(doc, "1.5 Secondary research questions", 2)
for item in [
    "Which responsibilities are centralised in one body, and which are split across advice, investigation, decision, discipline, prevention and publication?",
    "How do national systems allocate responsibility between institutions and national boards, offices, funders or sector bodies?",
    "Who can report, what conduct and populations fall within scope, and what thresholds govern admissibility?",
    "Which procedural safeguards are explicit for complainants, respondents, witnesses, experts and decision-makers?",
    "What information is published, with what denominator, at what level and after what form of anonymisation?",
    "How have the route, procedure and documentary framework changed over time?",
]:
    add_list_item(doc, item)

# 2 Methods
add_heading(doc, "2. Methods", 1)
add_heading(doc, "2.1 Study design", 2)
add_body(doc, "This is a descriptive comparative mapping study with a cross-sectional evidence cut-off and a historical component. The design combines documentary analysis, structured institutional case files, national system mapping and respondent validation. It is best understood as a multiple-case study in which universities are cases, route actors are nested units and national systems are contextual units.")
add_body(doc, "The analysis is explicitly non-evaluative. Differences in public output, procedure wording or institutional form are not interpreted as differences in integrity performance. The report distinguishes formal design, observed public reporting and member-confirmed practice.")

add_heading(doc, "2.2 Population and inclusion rule", 2)
add_body(doc, "The population consists of every university listed on the official LERU members page at the evidence cut-off of 18 August 2026: 24 universities in 12 countries [1]. No further inclusion or exclusion criteria are applied. UCLouvain is not included because it was not part of the official membership baseline used for the study. Membership in LERU-INTE or another integrity working group is a separate variable and remains unconfirmed for all 24 institutions.")

population_rows = []
for country, iso, institutions in COUNTRIES:
    population_rows.append((country, iso, len(institutions), "; ".join(institutions)))
add_caption(doc, "Table 1. Study population by country")
add_table(doc, ["Country", "ISO", "n", "LERU member universities"], population_rows, [1800, 650, 500, 6410], font_size=8.8)
add_source_note(doc, "Source: official LERU membership baseline [1]. Membership is the inclusion criterion; participation in a particular integrity group is not inferred from membership.")

add_heading(doc, "2.3 Units of analysis", 2)
add_body(doc, "Three nested levels are used:")
add_list_item(doc, "Inclusion unit: the official LERU member university.", bold_lead="Inclusion unit:")
add_list_item(doc, "Primary analytical unit: the institution's principal route for receiving, assessing, investigating or deciding on potential research-integrity breaches.", bold_lead="Primary analytical unit:")
add_list_item(doc, "Nested route actors: every committee, office, ombudsperson, adviser, referent, named person, panel, executive decision-maker, secretariat or publication owner that performs a distinct function.", bold_lead="Nested route actors:")
add_list_item(doc, "Contextual unit: the national research-integrity system, including law, codes, national boards, review routes, funders, academies and sector bodies.", bold_lead="Contextual unit:")

add_heading(doc, "2.4 Data-collection phases", 2)
phase_rows = [
    ("1", "Protocol and domain definition", "Nov 2024 onward", "Initial comparative objectives, committee questionnaire and procedural domains"),
    ("2", "LLM-assisted desk search", "2024–2025", "Multilingual country and institutional searches; structured deep-research outputs; source-link capture"),
    ("3", "Focused positioning analysis", "Jan 2026", "Response to the original request; preliminary institutional placement across all 24 members"),
    ("4", "Stakeholder refinement", "Jun 2026", "Feedback to make the report standalone and add descriptive numbers, comparisons and external comparators"),
    ("5", "Integration and provenance audit", "Aug 2026", "Per-institution dossiers, current-source retrieval, document manifests, data dictionary and evidence-gap audit"),
    ("6", "Targeted member validation", "Planned", "Institution-specific questions; correction of claims; acquisition of missing documents, composition and activity data"),
    ("7", "Final extraction and synthesis", "Planned", "Frozen comparative dataset, tables/figures, interpretive discussion and final report"),
]
add_caption(doc, "Table 2. Data-collection and integration phases")
add_table(doc, ["#", "Activity", "Timing", "Output"], phase_rows, [650, 2250, 1400, 5060], font_size=8.7)

add_heading(doc, "2.5 Information sources and evidence hierarchy", 2)
add_body(doc, "Evidence is prioritised in the following order, while recognising that different sources answer different questions:")
for item in [
    "Laws, ministerial instruments, national codes and official national-board rules.",
    "Current institutional statutes, regulations, complaints procedures and formal terms of reference.",
    "Official annual statements, statistics, published decisions, advice or case summaries.",
    "Official institutional webpages, organisational charts and public route descriptions.",
    "Direct written confirmation and documents supplied by institutional representatives.",
    "Secondary reports and earlier project summaries, used as discovery aids and clearly labelled when not independently verified.",
]:
    add_list_item(doc, item)
add_body(doc, "A direct email confirmation may be the best evidence for actual organisational placement, while a signed regulation remains the stronger source for formal authority. Conflicts are preserved and resolved explicitly rather than silently choosing one source.")

add_heading(doc, "2.6 Online search strategy and use of large language models", 2)
add_body(doc, "Desk research uses official-domain searches in English and the relevant national or institutional language. Core concepts include research/scientific integrity, misconduct, good research practice, commission/committee/board, ombudsperson, referent, reporting, complaint, investigation, annual report, case decision, appeal and the local equivalents. Searches proceed from the national system to the university route, the formal procedure, reports/statistics, related bodies and document history.")
add_body(doc, "Large language models are used to generate multilingual search terms, locate candidate pages, extract provisional fields, translate for screening, compare document structures and identify contradictions. They are not treated as sources. A model-generated statement can enter the report only after a human-auditable link to the original document or member confirmation has been recorded. Retrieval date, file checksum, version and pinpoint reference are stored where possible. Where a document is not machine readable, the extraction method is recorded.")

add_heading(doc, "2.7 Direct institutional validation", 2)
add_body(doc, "Direct feedback is used to confirm how written procedures operate, where the coordination function is administratively located, which bodies share responsibility and which documents or data are not public. The archive currently contains the original request, a preserved ETH response with two attachments, received Heidelberg documents without complete email provenance, meeting/context messages and oral feedback records. Relevant Outlook exports indicate additional attachments, but those attachments were not exported and cannot be treated as available evidence.")
add_body(doc, "The planned validation round is targeted rather than generic. Each institution will receive a short baseline summary and only the unresolved questions relevant to its dossier. Respondents will be asked to distinguish formal rules from current practice and to state whether their information may be quoted, paraphrased or used only for internal correction.")

add_heading(doc, "2.8 Data model: the comparative ‘family tree’", 2)
add_body(doc, "The data model is relational rather than a single narrative field. One university can contain multiple route actors; actors are linked to each other and to national bodies through directed relationships. Procedures, documents and events are versioned. Every substantive field carries a validation status and source reference.")

tree_lines = [
    "LERU population",
    "  └─ Country system (12)",
    "      ├─ law / national code / national oversight / review route",
    "      └─ LERU university (24)",
    "          ├─ organisational placement and governance",
    "          ├─ route actor(s): committee / office / ombuds / referent / panel",
    "          │   ├─ mandate, composition, appointment, independence",
    "          │   └─ relations to executives, legal, ethics, funders and national bodies",
    "          ├─ procedure version(s)",
    "          │   └─ intake → admissibility → inquiry → investigation → decision → review → publication",
    "          ├─ annual activity and outcomes",
    "          ├─ document and source manifest",
    "          ├─ dated events and document succession",
    "          └─ member validation and field-level provenance",
]
tree_table = doc.add_table(rows=1, cols=1)
tree_cell = tree_table.cell(0, 0)
tree_cell.text = ""
for idx, line in enumerate(tree_lines):
    p = tree_cell.paragraphs[0] if idx == 0 else tree_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    set_run_font(r, size=9.4, color=NAVY if idx == 0 else BLACK, bold=(idx == 0), name="Consolas")
configure_table(tree_table, [CONTENT_WIDTH_DXA], header=False, font_size=9.4)
set_cell_shading(tree_cell, LIGHT)
add_caption(doc, "Figure 1. Hierarchical study structure. Relations between actors are additionally stored as directed network edges.")

domain_rows = [
    ("1", "Population and identity", "Membership, aliases, country, INTE validation, inclusion provenance"),
    ("2", "National system", "Legal basis, codes, national bodies, autonomy, review and transparency"),
    ("3", "Route governance", "Body type, placement, mandate, powers, composition, appointment, independence"),
    ("4", "Procedure", "Scope, intake, admissibility, inquiry, investigation, evidence, timing, decision, appeal"),
    ("5", "Activity", "Signals, complaints, categories, durations, outcomes, measures and publication"),
    ("6", "Relationships", "Directed links among institutional, national, sector and funder actors"),
    ("7", "Documents", "Version, date, URL, local path, checksum, category, confidentiality and supported fields"),
    ("8", "Timeline", "LERU, national, institutional, route, procedure and document events"),
    ("9", "Email validation", "Claims confirmed/corrected, attachments, permission, open questions and response status"),
    ("10", "Field validation", "Evidence status, missingness, sources, confidence and resolution date"),
]
add_caption(doc, "Table 3. Core domains in the master data dictionary")
add_table(doc, ["#", "Object", "Core content"], domain_rows, [700, 2300, 6360], font_size=8.8)
add_source_note(doc, "The complete machine-readable dictionary is stored in 00_Shared_LERU_and_Methods/01_Study_Protocol_and_Codebook/MASTER_DATA_DICTIONARY_v0.1.json.")

add_heading(doc, "2.9 Document management, provenance and confidentiality", 2)
add_body(doc, "Each institutional dossier separates national context, institutional governance, procedures, reports/statistics/timeline, restricted correspondence, official web snapshots and working notes. A SHA-256 manifest records every local file. Publicly downloaded material retains its source URL and retrieval log. Correspondence is stored separately and excluded from any future public website unless redaction and permission have been checked.")
add_body(doc, "Source roles are distinguished: constitutive law or code; mandate; appointment/composition; complaints procedure; reporting form; annual report/statistics; decision/advice/case summary; monitoring/evaluation; training/prevention; organisational route; boundary-regime evidence; and correspondence/validation. Superseded documents are retained and linked to their replacements so that procedural change can be analysed.")

add_heading(doc, "2.10 Quality assurance and missing data", 2)
for item in [
    "Every claim receives one of: verified official public, verified member email, inferred, not found, not public, not applicable, conflicting sources or awaiting validation.",
    "‘Not found’ is never recoded as ‘absent’. Search coverage, source accessibility and language limitations are documented.",
    "Current and historical documents are separated; effective dates and supersession are recorded.",
    "Definitions and denominators are checked before numbers are placed in a shared table or graph.",
    "Research misconduct is kept separate from student assessment misconduct, research ethics review, clinical governance, HR grievances, fraud, data protection, IP, safety and security unless the local route explicitly connects them.",
    "A second extraction or targeted source check is required for high-impact variables such as decision authority, appeal rights, sanctions, anonymity and published case counts.",
]:
    add_list_item(doc, item)

add_heading(doc, "2.11 Descriptive analysis plan", 2)
add_body(doc, "All analyses are descriptive. Within each institution, the report will present the route map, procedure, documentary timeline and annual activity. Across institutions, categorical features will be summarised as n/N and percentages with explicit missing denominators. Comparable continuous variables—such as committee size, time limits or observed processing time—will be reported using median, interquartile range and full range. No inferential tests are planned because the population is purposively bounded and measurement heterogeneity is expected.")
analysis_rows = [
    ("Within institution", "Narrative case profile; actor-route diagram; procedure stages; document chronology; annual counts"),
    ("Within country", "National–institutional allocation; shared code/review route; differences among LERU members"),
    ("Across 24 members", "Frequencies, proportions, medians/ranges, cross-tabs and organisational/procedural typologies"),
    ("Across time", "LERU membership/integrity milestones; establishment and revision dates; annual activity trends"),
    ("Network", "Directed actor relationships: appointment, advice, investigation, decision, referral, review and publication"),
    ("Sensitivity", "Repeat summaries using only member-validated fields and only current primary documents"),
]
add_caption(doc, "Table 4. Planned descriptive analyses")
add_table(doc, ["Level", "Planned output"], analysis_rows, [2150, 7210], font_size=9.0)

add_heading(doc, "2.12 Ethical and data-protection considerations", 2)
add_body(doc, "The principal evidence consists of public governance documents and professional-role confirmations. Nevertheless, correspondence may contain personal contact information, views about sensitive cases or material supplied under an expectation of confidentiality. The report will not reproduce private email addresses, meeting links, identifiable allegations or quotations without permission. Case-level data will be limited to information officially published by the responsible body and will be reported at an aggregation level that respects the source's anonymisation.")

# 3 Baseline results
add_heading(doc, "3. Results: baseline population and evidence status", 1)
add_heading(doc, "3.1 LERU population and historical frame", 2)
add_body(doc, "The study population comprises 24 universities in 12 countries. The network began in 2002 with 12 members, expanded to 20 in 2006, added the University of Barcelona and Imperial College London in 2010, added the University of Copenhagen and Trinity College Dublin in 2017, and added ETH Zurich in 2024 [1,2]. These dates define the LERU membership frame; they are not the establishment dates of the universities' integrity functions.")
timeline_rows = [
    ("2002", "LERU founded with 12 member universities", "Network-level inclusion frame"),
    ("2006", "Expansion to 20 members", "Broader cross-national population"),
    ("2010", "Barcelona and Imperial join", "Membership reaches 22"),
    ("2017", "Copenhagen and Trinity join", "Membership reaches 23"),
    ("2020", "LERU publishes Towards a Research Integrity Culture at Universities", "Network-level integrity milestone"),
    ("2022", "Chair transition in the LERU Research Integrity Group recorded in project sources", "Group governance milestone; details to verify"),
    ("2024", "ETH Zurich joins; LERU issues Communicating with Integrity", "Membership reaches 24; integrity-related policy milestone"),
    ("2026", "Standalone comparative report and institutional dossiers initiated", "Current project integration phase"),
]
add_caption(doc, "Table 5. Initial LERU and project timeline")
add_table(doc, ["Year", "Event", "Analytical relevance"], timeline_rows, [1000, 4450, 3910], font_size=8.8)

add_heading(doc, "3.2 Evidence coverage before normalised extraction", 2)
add_body(doc, "The latest accessible structured LERU baseline contains a public institutional procedure or route for 24/24 members. Fourteen profiles are coded as strong and ten as moderate evidence. A strict public-output indicator is positive for 14/24. The source index contains 97 institutional links: 25 direct PDFs, one direct Word file, two JSON table endpoints and 69 HTML or other web pages. Exact matching with the earlier formal source registry succeeds for only 29 of the 97 links, indicating provenance drift that must be repaired before final analysis.")
add_body(doc, "The dossier-building pass preserved the legacy evidence, copied the available received documents, performed an automated retrieval of direct legacy links and then added a curated current-source pass. The first pass produced 19 successful and eight failed retrievals; the curated pass produced 20 successful and three failed retrievals. These are retrieval events, not counts of unique or fully validated documents. Failures and access restrictions remain visible in the logs.")

trans_rows = [
    ("Institution-owned local output", "13", "Local annual report, statement, cases, minutes or statistics found"),
    ("National/sector output", "4", "Relevant output is primarily published outside the institution"),
    ("Procedure only", "5", "A route/procedure is public, but no qualifying output was found"),
    ("Restricted/internal output", "1", "Output exists or is referenced but is not openly accessible"),
    ("Historical/case-specific", "1", "Only historical or isolated case-level output met the broader criterion"),
]
add_caption(doc, "Table 6. Preliminary public-output typology (n = 24)")
add_table(doc, ["Category", "n", "Interpretation"], trans_rows, [2800, 700, 5860], font_size=9.0)
add_source_note(doc, "These categories measure what could be evidenced publicly in the current source baseline. They do not measure the quality, fairness or effectiveness of the underlying systems.")

add_heading(doc, "3.3 Current dossier coverage", 2)
coverage_rows = []
for country, iso, institutions in COUNTRIES:
    for institution in institutions:
        cov = coverage[institution]
        coverage_rows.append((
            iso,
            institution,
            cov["local_file_count"],
            cov["procedure_or_policy_files"],
            cov["report_statistics_timeline_files"],
            cov["restricted_correspondence_files"],
        ))
add_caption(doc, "Table 7. Local dossier file coverage after source migration and current-source retrieval")
add_table(doc, ["Country", "Institution", "All indexed files", "Procedure/policy", "Reports/data", "Restricted mail"], coverage_rows, [850, 2870, 1250, 1450, 1400, 1540], font_size=8.0)
add_source_note(doc, "Generated from the SHA-256 dossier manifests. Counts include structured seeds and metadata but exclude the manifest itself. Zero means no separate local file in that subfolder, not that the information or route is absent.")

add_heading(doc, "3.4 Preliminary answer to the original positioning question", 2)
add_body(doc, "The first-pass synthesis suggests four recurring placement patterns: (1) a coordination or research-integrity unit inside a research office; (2) a university-wide committee appointed by the executive and supported by Legal Affairs or governance; (3) a central executive route through a Rectorate, Chancellor or Registrar; and (4) a distributed model in which advice, investigation, decision, discipline and prevention sit in different units. These patterns are analytically useful but not yet counted because the same institution can legitimately occupy more than one category.")
add_callout(doc, "Key interpretive distinction", "Administrative placement is not the same as investigative independence. A committee can be supported by Legal Affairs without Legal Affairs determining its findings; conversely, a research-office location does not by itself establish that the office investigates or decides cases. The final dataset therefore codes host unit, secretariat, investigators, adviser, decision-maker and appeal/review body separately.")

# 4 Country chapters
add_heading(doc, "4. Results: country and institution profiles", 1)
add_body(doc, "The following profiles are structured as a result-section shell with a preliminary evidence baseline. They provide enough information to guide extraction and targeted validation, but they are not yet final institutional reports. Each profile will ultimately contain a national-system narrative, actor map, procedure, document chronology, caseload/outcomes, interfaces and a fully sourced evidence table.")

country_number = 0
for country, iso, institutions in COUNTRIES:
    country_number += 1
    doc.add_page_break()
    add_heading(doc, f"4.{country_number} {country}", 2)
    add_callout(doc, "National system — working classification", COUNTRY_SYSTEMS[country], fill=PALE_BLUE)
    add_body(doc, f"LERU institutions in scope: {', '.join(institutions)}. The final country section will specify the legal and normative hierarchy, national actors, institutional autonomy, review/appeal routes, transparency obligations, funder interfaces and the boundary with ethics, disciplinary law, student misconduct and other adjacent regimes.")
    for institution in institutions:
        add_profile(doc, institution, coverage)

# 5 Synthesis
add_heading(doc, "5. Cross-case synthesis and planned comparative outputs", 1)
add_heading(doc, "5.1 Initial structural patterns", 2)
add_body(doc, "The baseline already shows that a single label such as ‘scientific integrity committee’ obscures important differences. At minimum, the following structural patterns require separate coding:")
for item in [
    "Executive-centred routes, in which the Rector, Chancellor or Registrar is the named institutional authority.",
    "Committee-to-executive routes, in which an independent or semi-independent committee investigates and advises a governing board that decides.",
    "Research-office coordination routes, in which a dedicated office manages intake, policy, training, files and committee support.",
    "Named-officer or referent routes, in which one institutional officer receives and investigates reports, sometimes with experts or an ad-hoc panel.",
    "National–institutional two-track routes, in which legally defined misconduct is examined nationally and other deviations remain local.",
    "Distributed hybrids, in which prevention, case handling, discipline, student matters, ethics and publication are deliberately separated.",
]:
    add_list_item(doc, item)

add_heading(doc, "5.2 Dimensions that can be compared safely", 2)
comparison_rows = [
    ("Organisational", "Host unit; secretariat; reporting line; permanence; role separation"),
    ("Governance", "Mandate; composition; appointment; term; external members; conflicts; independence safeguards"),
    ("Access", "Who may report; public/internal channels; anonymity; support; retaliation protection"),
    ("Scope", "Populations; research stages; FFP; other practices; time limits; cross-institution/funder cases"),
    ("Procedure", "Triage; admissibility; preliminary inquiry; investigation; evidence; hearings; timelines"),
    ("Outcome", "Finding categories; decision authority; measures; referrals; review/appeal"),
    ("Transparency", "Annual statements; counts; decisions; anonymisation; denominators; publication owner"),
    ("Prevention", "Advisers; training; guidance; policy development; learning from cases"),
    ("History", "Establishment; predecessor bodies; procedure revisions; major reports and network milestones"),
]
add_caption(doc, "Table 8. Core comparative dimensions")
add_table(doc, ["Dimension", "Variables"], comparison_rows, [2100, 7260], font_size=9.0)

add_heading(doc, "5.3 Planned tables and figures", 2)
for item in [
    "Table A: population, route type, evidence status and validation status for all 24 institutions.",
    "Table B: national system typology, responsible national bodies and institutional/national division of labour.",
    "Table C: institutional placement, secretariat, investigators, adviser and final decision-maker.",
    "Table D: composition, appointment, term, external expertise and conflict-of-interest safeguards.",
    "Table E: procedural comparison from intake through review, including scope and time limits.",
    "Table F: five-year complaint/case counts, allegation categories, outcomes and denominators where comparable.",
    "Table G: public-output and verification matrix, including not-found/not-public distinctions.",
    "Figure A: LERU membership and research-integrity milestone timeline.",
    "Figure B: complaint-workflow archetypes.",
    "Figure C: heatmap of data availability and public transparency variables.",
    "Figure D: network of institutional and national bodies, limited to verified relationships.",
    "Figure E: small-multiple timelines for committee/procedure establishment and revision.",
]:
    add_list_item(doc, item)

add_heading(doc, "5.4 Rules for quantitative summaries", 2)
add_body(doc, "Counts will be reported only when the event definition and observation window are explicit. ‘Signals’, ‘formal complaints’, ‘admissible complaints’, ‘investigations’ and ‘decisions’ are not interchangeable. A rate will not be calculated without a defensible denominator. If institutions publish only small counts or narrative cases, these will remain descriptive. Missing years will not be treated as zero. Where five-year series are sufficiently complete, the report will show institution-specific trends and pooled descriptive totals with the number of contributing institutions for every year.")

# 6 Discussion
add_heading(doc, "6. Discussion", 1)
add_heading(doc, "6.1 Preliminary interpretation", 2)
add_body(doc, "The most important early finding is conceptual: LERU universities do not share one uniform organisational object called a research-integrity committee. They share a responsibility that is distributed across different actors. Comparative validity therefore depends on following the function rather than matching names. National systems help explain some of the variation—for example, whether legally defined misconduct is investigated by a national body—but do not fully determine local governance, support or transparency.")
add_body(doc, "A second early finding is that organisational placement, investigative independence and decision authority must be analysed separately. The office that receives a report may not investigate it; the committee that investigates may only advise; the executive that decides may rely on Legal Affairs for process support; and publication may sit with yet another unit. Treating the visible contact page as ‘the committee’ would collapse these differences.")
add_body(doc, "A third finding is that public transparency is structurally heterogeneous. Some institutions publish detailed annual statements or anonymised cases, whereas others rely on national repositories, sector reports or procedures alone. This affects what can be observed and compared. It does not justify conclusions about underlying incidence or institutional performance without validated denominators and direct confirmation.")

add_heading(doc, "6.2 Strengths", 2)
for item in [
    "A complete and externally verifiable population rather than a convenience sample.",
    "A multilevel design linking country systems, institutional routes, individual actors, documents and time.",
    "Triangulation of primary public sources, preserved received documents and targeted member feedback.",
    "Explicit separation of public evidence, restricted correspondence, inference and missingness.",
    "Versioned source manifests and a machine-readable data dictionary that can support later reproducibility and website publication.",
    "A function-based definition that accommodates committees, offices, ombudspersons, referents and hybrid systems without forcing false equivalence.",
]:
    add_list_item(doc, item)

add_heading(doc, "6.3 Limitations", 2)
for item in [
    "The current baseline is not yet a completed normalised comparative dataset; many details remain in narrative or heterogeneous legacy schemas.",
    "LERU-INTE participation has not been confirmed for the 24 included universities.",
    "Public sources may be outdated, incomplete, language-dependent or silent on internal practice, staffing and informal triage.",
    "Only 29 of 97 institutional source links match the earlier formal source registry exactly; provenance and URL variants need reconciliation.",
    "Committee establishment dates, composition history, processing time, complaint categories, outcomes and measures are frequently missing.",
    "The live Outlook mailbox and non-exported attachments were not accessible; the correspondence base is therefore incomplete.",
    "Comparisons of case counts are vulnerable to different definitions, reporting thresholds, observation periods and small-number suppression.",
    "LLM-assisted extraction can introduce translation or interpretation errors; the workflow controls but cannot eliminate this risk until source-level verification is complete.",
]:
    add_list_item(doc, item)

add_heading(doc, "6.4 Implications for LERU and participating universities", 2)
add_body(doc, "A practical LERU comparison should adopt a minimum common data standard rather than require identical procedures. That standard could define a small set of stable variables: route actors and roles; administrative host and secretariat; scope and intake; investigation/decision/review allocation; composition and independence safeguards; procedure version; annual signals/complaints/investigations/outcomes with definitions; and public-output location. A shared standard would improve mutual learning without prescribing one governance model.")
add_body(doc, "For institutional design, the comparison can clarify trade-offs between research-office proximity, legal/process expertise, executive authority and perceived independence. It can also reveal where responsibilities are duplicated or where complainants may encounter a fragmented route. These implications should be discussed only after member validation, particularly when an organisational placement is under active review.")

add_heading(doc, "6.5 Priorities before finalisation", 2)
for item in [
    "Complete the same core field set for every institution before drawing comparative conclusions.",
    "Run a source-currency check and link superseded documents to current versions.",
    "Send targeted member summaries and questions; record confirmations, corrections and permission.",
    "Extract annual reports into a denominator-aware activity table and flag incomparable series.",
    "Build verified actor relationships and timelines instead of deriving history from free-text year mentions.",
    "Freeze a versioned analysis dataset before producing final tables, figures or a public website layer.",
]:
    add_list_item(doc, item)

# 7 Conclusion
add_heading(doc, "7. Conclusion", 1)
add_body(doc, "The project now has a coherent scientific architecture for answering the original question and extending it into a rigorous comparison of research-integrity functions across LERU. The preliminary evidence confirms a relevant route for every member university and reveals substantial variation in national context, institutional placement, role allocation and public output. The decisive next step is not further broad searching alone, but structured extraction and institution-specific validation using the common data model. Until that step is complete, the report should present system architecture and evidence gaps transparently and avoid rankings or unsupported numerical comparisons.")

# References
add_heading(doc, "References and core source framework", 1)
references = [
    ("League of European Research Universities. Members.", "https://www.leru.org/members", "Accessed 18 August 2026."),
    ("League of European Research Universities. The First Decade, 2002–2012.", "https://www.leru.org/files/LERU10LOWRES_final.pdf", None),
    ("LERU. Towards a Research Integrity Culture at Universities: From Recommendations to Implementation (2020).", "https://www.leru.org/files/Towards-a-Research-Integrity-Culture-at-Universities-full-paper.pdf", None),
    ("LERU. Communicating with Integrity (2024).", "https://www.leru.org/files/Publications/Communicating-with-intergrity_LERU-paper.pdf", None),
    ("ALLEA. The European Code of Conduct for Research Integrity, revised edition (2023).", "https://allea.org/code-of-conduct/", None),
    ("Finnish National Board on Research Integrity TENK. The Finnish code of conduct for research integrity and procedures for handling alleged violations (2023).", "https://tenk.fi/sites/default/files/2023-11/RI_Guidelines_2023.pdf", None),
    ("KU Leuven. Regulation/Procedure for handling complaints of alleged violations of research integrity.", "https://research.kuleuven.be/en/integrity-ethics/integrity/procedures", "Current official page checked 18 August 2026."),
    ("University of Amsterdam. Academic Integrity Complaints Regulations and CWI annual report 2024.", "https://www.uva.nl/en/research/research-environment/academic-integrity/submitting-a-complaint-about-a-breach-of-academic-integrity/submitting-a-complaint-about-a-breach-of-academic-integrity.html", None),
    ("Lund University. Deviations from good research practice; institutional guidelines and annual reports.", "https://www.staff.lu.se/research-and-education/research-support/research-ethics-and-animal-testing-ethics/deviations-good-research-practice", None),
    ("Imperial College London. Regulation 21 — Research Misconduct Policy and Procedures, February 2026.", "https://www.imperial.ac.uk/media/imperial-college/administration-and-support-services/secretariat/new-library-may-24/Regulation-21---Research-Misconduct-Policy-and-Procedures.pdf", None),
    ("University of Oxford. Procedure for Addressing Potential Breaches of Research Integrity.", "https://www.ox.ac.uk/research/support/governance-and-committees/research-policies/procedure-for-addressing-potential", "Official current procedure page checked 18 August 2026."),
    ("Project archive. European committees draft protocol; research integrity systems codebook; original positioning request; institutional positioning report; received documents and meeting records.", None, "Restricted and public local project sources, as applicable."),
    ("Project master data dictionary v0.1 and dossier manifests.", None, "Generated 18 August 2026; working methodological instruments."),
]
for idx, (title_text, url, note) in enumerate(references, start=1):
    add_reference(doc, idx, title_text, url, note)

# Appendices
doc.add_page_break()
add_heading(doc, "Appendix A. Full institutional dossier template", 1)
add_body(doc, "Each institution chapter will be completed in the following order. The template is deliberately broader than the original positioning question so that a new institution could be added without redesigning the study.")
appendix_a = [
    ("A1", "Identity and inclusion", "Official/English name; aliases; country; city; LERU status/date/source; INTE confirmation"),
    ("A2", "National system", "Law and codes; responsible national actors; institutional autonomy; review/appeal; publication; boundary regimes"),
    ("A3", "Institutional placement", "Host unit; organisational chart; reporting line; secretariat; staffing/resources; legal involvement"),
    ("A4", "Route actors", "Official names; type; mandate; authority; decision-maker; adviser; investigators; publication owner"),
    ("A5", "Composition and independence", "Number/roles; expertise; external members; appointment; term; recusal; conflicts; safeguards"),
    ("A6", "Scope and access", "Who may report; respondent/research scope; FFP/QRP/other; anonymity; limitation period; support/protection"),
    ("A7", "Procedure", "Intake; admissibility; preliminary inquiry; investigation; evidence; hearings; timing; outcome; measures; review"),
    ("A8", "Related bodies", "Executive, Legal, HR, ethics, student, privacy, IP, safety, funders, national boards, publishers and partners"),
    ("A9", "Activity and transparency", "Signals; complaints; investigations; durations; allegation types; outcomes; measures; reports/cases; denominators"),
    ("A10", "Documents", "Constitutive, procedure, form, annual report, case output, evaluation, training, organisation and correspondence"),
    ("A11", "Timeline", "Establishment; predecessor; appointments; revisions; major cases/reports; national and LERU milestones"),
    ("A12", "Validation", "Field-level status, conflicts, missingness, targeted questions, respondent corrections and permission"),
]
add_table(doc, ["Block", "Section", "Required content"], appendix_a, [800, 2400, 6160], font_size=8.7)

add_heading(doc, "Appendix B. Blank source and document inventory", 1)
doc_rows = [
    ("DOC-…", "", "", "", "", "", "", "", ""),
    ("DOC-…", "", "", "", "", "", "", "", ""),
    ("DOC-…", "", "", "", "", "", "", "", ""),
]
add_table(doc, ["ID", "Title", "Category", "Publisher", "Version/date", "Language", "URL/path", "SHA-256", "Fields supported"], doc_rows, [700, 1500, 1100, 1100, 900, 700, 1500, 900, 960], font_size=7.4)
add_source_note(doc, "Add one row per document or web snapshot. Restricted correspondence must use a project reference rather than personal details in any public export.")

add_heading(doc, "Appendix C. Targeted member-validation question bank", 1)
add_body(doc, "The final email to each member should include only unanswered questions. A short baseline summary should be offered for correction before requesting additional data.")
questions = [
    ("Core", "What is the official name of the function(s) that receive, assess, investigate and decide on research-integrity concerns?"),
    ("Core", "Where is each function administratively located, who provides the secretariat, and to whom does it report?"),
    ("Core", "Which body/person investigates, which advises, which decides and which can impose or initiate measures?"),
    ("Core", "Please confirm the current governing regulation/procedure and its effective date; are any revisions planned?"),
    ("Core", "What populations, research activities and categories of conduct fall inside or outside the procedure?"),
    ("Core", "Who can report, are anonymous reports possible, and which admissibility or limitation rules apply?"),
    ("Core", "What are the stages and formal or usual timeframes from receipt to final decision?"),
    ("Core", "How are committee/panel members appointed, for how long, and how are expertise, external membership and conflicts handled?"),
    ("Core", "What internal review, national second-opinion, appeal or judicial routes are available?"),
    ("Core", "Which other university, national, funder or partner bodies must be informed or involved?"),
    ("Data", "For each of the last five reporting years, how many signals, formal complaints, admissible complaints, investigations and decisions were recorded?"),
    ("Data", "Can counts be divided by allegation category, outcome, measure and review/appeal, with the definitions and denominators used?"),
    ("Data", "Are median/typical processing time and backlog available, even if they cannot be published?"),
    ("Documents", "Please provide or link the current mandate/terms of reference, procedure, reporting form, composition decision, annual reports and public cases."),
    ("History", "When was the current function established, what preceded it, and when were the major procedural revisions?"),
    ("Permission", "May we quote or paraphrase this confirmation and identify your institutional role, or should it be used only to correct the public-source summary?"),
]
add_table(doc, ["Type", "Question"], questions, [1100, 8260], font_size=8.8)

add_heading(doc, "Appendix D. Minimum comparable dataset for a new institution", 1)
add_body(doc, "A new institution can be added when the following minimum is available: verified membership/identity; national system summary; at least one primary route actor; current procedure; investigator/adviser/decision allocation; scope and intake; review route; one source-manifest entry per claim; transparency status; and an explicit validation/missingness record. Composition, activity and historical fields may remain missing, but the missingness must be coded rather than left implicit.")

add_heading(doc, "Appendix E. Evidence-status legend", 1)
legend_rows = [
    ("Verified — official public", "Directly supported by a current official public source"),
    ("Verified — member email", "Confirmed by an attributable institutional representative; publication permission recorded separately"),
    ("Inferred", "Reasonable synthesis from sources, but not explicitly stated"),
    ("Not found", "Search completed to the recorded extent; absence is not established"),
    ("Not public", "Source or information is known/referenced but not publicly accessible"),
    ("Not applicable", "The field does not apply to the local structure, with reason"),
    ("Conflicting sources", "Sources disagree or describe different time periods/levels"),
    ("Awaiting validation", "Provisional statement prepared for targeted institutional confirmation"),
]
add_table(doc, ["Status", "Meaning"], legend_rows, [2600, 6760], font_size=9.0)

add_callout(doc, "End-of-draft status", "This document is a complete scientific report framework and preliminary evidence baseline. It is intentionally explicit about what is not yet validated. The next version should replace provisional profile statements with field-level sourced results and populate the planned comparative tables and figures from the frozen dataset.", fill=PALE_GOLD, title_color=GOLD)

# Keep final paragraph for stable end marker.
endp = doc.add_paragraph()
endp.alignment = WD_ALIGN_PARAGRAPH.CENTER
endp.paragraph_format.space_before = Pt(12)
er = endp.add_run("— End of working draft —")
set_run_font(er, size=8.8, color=MUTED, italic=True)

# Global section settings and compatibility.
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

settings = doc.settings.element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.save(OUTPUT)
print(f"Created {OUTPUT}")
